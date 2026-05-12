"""DOCX 솔루션 계획서 검수기 v2.

SolutionPlan v2 ↔ 생성된 DOCX 파일을 교차 검증합니다.
양식: ValueChain 표준 (data/references/_솔루션양식_표준.md)

검수 항목:
1. 핵심 지표 박스 (2~6건) 존재
2. 핵심 지표 값 JSON 일치
3. 비교표 구조 및 카테고리
4. 본문 섹션 전체 포함
5. 항목 번호 (직렬/섹션별) 모드 일관성
6. 사장님 확인 섹션 (옵션)
7. 수수료 구조 (옵션)
8. 업장명 일치
9. 회사명 푸터 (옵션)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from src.schemas.solution import SolutionPlan
from src.schemas.validation import CheckGroup, CheckItem, CheckStatus


def _extract_all_text(doc: Document) -> str:
    """문서 전체 텍스트 추출 (테이블 포함)."""
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def validate_docx(plan: SolutionPlan, docx_path: str | Path) -> CheckGroup:
    """SolutionPlan v2 ↔ DOCX 교차 검증."""
    docx_path = Path(docx_path)
    items: list[CheckItem] = []

    if not docx_path.exists():
        items.append(CheckItem(
            name="파일 존재", status=CheckStatus.FAIL,
            message=f"{docx_path.name} 파일을 찾을 수 없습니다",
        ))
        return CheckGroup(name="DOCX 검수", items=items)

    doc = Document(str(docx_path))
    all_text = _extract_all_text(doc)

    # ── 1. 핵심 지표 박스 (T0이 KPI 테이블) ──
    expected_n = len(plan.core_metrics)
    kpi_table = doc.tables[0] if doc.tables else None
    actual_n = len(kpi_table.columns) if kpi_table else 0
    items.append(CheckItem(
        name=f"핵심 지표 박스 {expected_n}건",
        status=CheckStatus.PASS if actual_n == expected_n else CheckStatus.FAIL,
        message=f"기대 {expected_n}건 / 실제 {actual_n}건",
        expected=str(expected_n),
        actual=str(actual_n),
    ))

    # ── 2. 핵심 지표 값 JSON 일치 ──
    kpi_values_in_doc: list[str] = []
    if kpi_table:
        for cell in kpi_table.rows[0].cells:
            kpi_values_in_doc.append(cell.text.strip())

    # 업주 docx 표시 변환: is_target 인 KPI 는 docx 에서 만원 반올림된 값으로 노출됨.
    # 검수 시에도 원본 값 OR 반올림값 둘 중 하나와 매치되면 통과.
    def _round_to_10k(amount: int) -> int:
        if amount == 0:
            return 0
        sign = -1 if amount < 0 else 1
        return sign * ((abs(amount) + 5_000) // 10_000) * 10_000

    def _rounded_won_str(raw_value: str) -> str | None:
        digits = "".join(ch for ch in raw_value if ch.isdigit())
        if not digits:
            return None
        try:
            val = int(digits)
        except ValueError:
            return None
        suffix = "원" if "원" in raw_value else ""
        return f"{_round_to_10k(val):,}{suffix}"

    # 50만원 내림 헬퍼 (목표 KPI 매칭용 — 2026-04-22 피드백)
    def _round_down_to_500k_local(amount: int) -> int:
        if amount < 100_000:
            return amount
        return (amount // 500_000) * 500_000

    def _floored_won_str_local(raw_value: str) -> str | None:
        digits = "".join(ch for ch in raw_value if ch.isdigit())
        if not digits:
            return None
        try:
            val = int(digits)
        except ValueError:
            return None
        suffix = "원" if "원" in raw_value else ""
        return f"{_round_down_to_500k_local(val):,}{suffix}"

    mismatches: list[str] = []
    for metric in plan.core_metrics:
        candidates = [metric.value]
        # baseline KPI: docx 노출 시 만원 반올림
        # target KPI: docx 노출 시 50만원 단위 내림 (2026-04-22 피드백)
        if metric.is_baseline:
            rounded = _rounded_won_str(metric.value)
            if rounded and rounded != metric.value:
                candidates.append(rounded)
        if metric.is_target:
            rounded = _rounded_won_str(metric.value)
            if rounded and rounded != metric.value:
                candidates.append(rounded)
            floored = _floored_won_str_local(metric.value)
            if floored and floored != metric.value:
                candidates.append(floored)
        found = any(
            any(c in cell_text for cell_text in kpi_values_in_doc)
            for c in candidates
        )
        if not found:
            mismatches.append(f"{metric.label}={metric.value}")

    items.append(CheckItem(
        name="핵심 지표 값 일치",
        status=CheckStatus.PASS if not mismatches else CheckStatus.WARN,
        message=(
            "전체 핵심 지표 값 일치"
            if not mismatches
            else f"불일치: {', '.join(mismatches[:3])}"
        ),
    ))

    # ── 3. 비교표 구조 및 카테고리 ──
    comp_table = doc.tables[1] if len(doc.tables) > 1 else None
    expected_rows = len(plan.comparison.rows)
    if comp_table and len(comp_table.rows) >= expected_rows + 1:
        comp_ok = True
        comp_issues: list[str] = []

        # 헤더 라벨
        headers = " ".join(c.text.strip() for c in comp_table.rows[0].cells)
        expected_label = plan.comparison.header_label
        # "Before/After" 또는 "지금/앞으로"가 들어있는지
        if expected_label == "지금/앞으로":
            if "지금" not in headers or "앞으로" not in headers:
                comp_ok = False
                comp_issues.append("헤더에 '지금/앞으로' 없음")
        else:
            if "Before" not in headers or "After" not in headers:
                comp_ok = False
                comp_issues.append("헤더에 'Before/After' 없음")

        # 카테고리 일치
        for ri, comp_row in enumerate(plan.comparison.rows):
            row_text = " ".join(
                c.text.strip() for c in comp_table.rows[ri + 1].cells
            )
            if comp_row.category not in row_text:
                comp_ok = False
                comp_issues.append(f"'{comp_row.category}' 행 없음")

        items.append(CheckItem(
            name="비교표 구조·카테고리",
            status=CheckStatus.PASS if comp_ok else CheckStatus.FAIL,
            message=(
                f"비교표 {expected_rows}행 정상"
                if comp_ok
                else f"이슈: {', '.join(comp_issues[:3])}"
            ),
        ))
    else:
        items.append(CheckItem(
            name="비교표 구조·카테고리",
            status=CheckStatus.FAIL,
            message=f"비교표 {expected_rows}행 미발견",
        ))

    # ── 4. 본문 섹션 제목 전체 포함 ──
    section_titles = [s.title for s in plan.sections]
    missing_sections = [t for t in section_titles if t not in all_text]
    items.append(CheckItem(
        name="섹션 제목 전체 포함",
        status=CheckStatus.PASS if not missing_sections else CheckStatus.FAIL,
        message=(
            f"전체 {len(section_titles)}개 섹션 존재"
            if not missing_sections
            else f"누락: {', '.join(missing_sections[:3])}"
        ),
    ))

    # ── 5. 항목 번호 중복 검사 ──
    # 완화 정책 (2026-04-26): 담당자 수동 편집·섹션 번호 재시작 등으로 번호가
    # 비연속적일 수 있다. 엄격한 연속성 대신 "심각한 중복(3회 이상 반복 or
    # 중복 번호가 전체 고유값의 과반)"만 WARN. 단순 재시작 수준은 PASS.
    item_numbers: list[str] = []
    for p in doc.paragraphs:
        m = re.match(r"^([\d]+(?:\.[\d]+)?)\.\s", p.text.strip())
        if m:
            item_numbers.append(m.group(1))

    numbering_ok = True
    duplicate_msg = ""
    seen: dict[str, int] = {}
    if plan.item_numbering == "serial":
        top_level = [n for n in item_numbers if "." not in n]
        for n in top_level:
            seen[n] = seen.get(n, 0) + 1
        # 동일 번호가 3회 이상이거나, 중복 번호가 과반이면 의심
        heavy = [n for n, c in seen.items() if c >= 3]
        dup = [n for n, c in seen.items() if c > 1]
        dup_ratio = len(dup) / len(seen) if seen else 0.0
        if heavy or dup_ratio > 0.5:
            numbering_ok = False
            duplicate_msg = (
                f"3회 이상 반복: {', '.join(heavy[:3])}"
                if heavy
                else f"중복 번호 비율 {dup_ratio:.0%}: {', '.join(dup[:3])}"
            )
    # PASS 메시지에 참고 중복 표기 (있으면)
    note_parts = [f"{len(item_numbers)}개 항목 번호"]
    if seen:
        dup_note = [n for n, c in seen.items() if c > 1]
        if dup_note:
            note_parts.append(f"섹션 재시작 참고: {', '.join(dup_note[:3])}")
        else:
            note_parts.append("중복 없음")
    pass_msg = " — ".join(note_parts)

    items.append(CheckItem(
        name=f"항목 번호 ({plan.item_numbering})",
        status=CheckStatus.PASS if numbering_ok else CheckStatus.WARN,
        message=(
            pass_msg if numbering_ok
            else f"항목 번호 심각한 중복 — {duplicate_msg}"
        ),
    ))

    # ── 6. 사장님 확인 (옵션) ──
    if plan.owner_checks:
        owner_check_text = "사장님 확인 필요 사항"
        found = owner_check_text in all_text
        items.append(CheckItem(
            name="사장님 확인 섹션",
            status=CheckStatus.PASS if found else CheckStatus.FAIL,
            message=(
                f"⑥ {owner_check_text} 섹션 존재 ({len(plan.owner_checks)}건)"
                if found
                else "⑥ 사장님 확인 섹션 누락"
            ),
        ))

    # ── 7. 수수료 구조 (옵션) ──
    # 업주 docx (2026-04-22 피드백): 수수료 구간 금액은 50만원 단위 내림 표시.
    # 원본 값 OR 만원 반올림값 OR 50만원 내림값 중 하나라도 일치하면 통과.
    if plan.fee_structure:
        fee_found = "수수료" in all_text

        def _round_down_to_500k(amount: int) -> int:
            if amount < 100_000:
                return amount
            return (amount // 500_000) * 500_000

        def _floored_won_str(raw_value: str) -> str | None:
            digits = "".join(ch for ch in raw_value if ch.isdigit())
            if not digits:
                return None
            try:
                val = int(digits)
            except ValueError:
                return None
            suffix = "원" if "원" in raw_value else ""
            return f"{_round_down_to_500k(val):,}{suffix}"

        def _tier_amount_match(raw: str) -> bool:
            if raw in all_text:
                return True
            rounded = _rounded_won_str(raw)
            if rounded and rounded in all_text:
                return True
            floored = _floored_won_str(raw)
            if floored and floored in all_text:
                return True
            return False

        amount_match = all(
            _tier_amount_match(tier.amount) for tier in plan.fee_structure.tiers
        )
        items.append(CheckItem(
            name="수수료 구조 존재 + 금액 일치",
            status=CheckStatus.PASS if (fee_found and amount_match) else CheckStatus.FAIL,
            message=(
                f"⑦ 수수료 ({len(plan.fee_structure.tiers)}구간) 정상"
                if fee_found and amount_match
                else "수수료 섹션 또는 금액 누락"
            ),
        ))

    # ── 7.5. 핵심 메시지 박스 체크 제거 ──
    # 담당자 피드백(2026-04-22): 업주 제공용 docx 에서 핵심 메시지 박스 제외.
    # solution_plan.json 에는 key_message 필드가 유지되지만 docx 에는 노출하지 않으므로
    # 이 체크 자체를 제거 (검수 로직 ↔ 빌더 출력 동기화).

    # ── 8. 업장명 일치 ──
    items.append(CheckItem(
        name="업장명 일치",
        status=CheckStatus.PASS if plan.store.name in all_text else CheckStatus.FAIL,
        message=(
            f"'{plan.store.name}' 확인됨"
            if plan.store.name in all_text
            else f"'{plan.store.name}' 누락"
        ),
    ))

    # ── 9. 회사명 푸터 ──
    items.append(CheckItem(
        name="회사명 푸터",
        status=CheckStatus.PASS if plan.company_name in all_text else CheckStatus.WARN,
        message=(
            f"'{plan.company_name}' 푸터 확인됨"
            if plan.company_name in all_text
            else f"'{plan.company_name}' 푸터 누락"
        ),
    ))

    return CheckGroup(name="DOCX 검수", items=items)
