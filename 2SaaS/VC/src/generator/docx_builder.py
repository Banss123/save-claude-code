"""솔루션 계획서 DOCX 생성기 v2.

ValueChain 표준 양식 (5건 검증 후 확정)에 맞춘 결정적 빌더.
양식 명세 마스터: data/references/_솔루션양식_표준.md

핵심 설계:
- 가변 KPI 박스 (2~6개)
- 가변 비교표 행 (3~5행, 멀티라인)
- 가변 본문 섹션 (3~8개)
- 마커 분기 (★/▸/빨간경고)
- 회색/녹색 인용 박스
- 꼬리 (사장님 확인 + 수수료) 옵션

동일 입력 = 동일 출력 보장.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.schemas.solution import Section, SectionItem, SolutionPlan

# ── L-2 정본 목표매출 산정서 상수 ──
# 정본 §출력 형식 (라인 199~303) 그대로 반영.
# lever_report 존재 시 기존 tier 표 + 근거 박스를 대체해 6개 섹션을 렌더한다.
_SECTION_RULE_CHAR = "━"
_SECTION_RULE_WIDTH = 37  # 정본 문서의 구분선 폭

# 매장주 실행 필수 체크 항목 (정본 §4)
_MANDATORY_OWNER_CHECKS: list[str] = [
    "메뉴 사진 교체 (대표 메뉴 최우선)",
    "리뷰 24시간 내 답변 유지",
    "영업시간/휴무일 정합성 유지",
]

# 위험(경고) 배경색 — 수수료 상한 재조정 / 희망매출 괴리 시
WARNING_BG = "FFF4E5"
WARNING_BORDER = "E68A00"

# 월 표시 → "M월차 (YYYY-MM 월)" 같은 세로 서식 대신, 담당자 친숙하게 "N개월 후" 유지.
# tier_plan 의 month 필드는 "목표 도달 월" 자체로 노출하지 않고, 3/6/12 개월 라벨만 쓴다.
_TIER_LABELS: dict[str, str] = {
    "tier1_3m": "3개월차",
    "tier2_6m": "6개월차",
    "tier3_12m": "12개월차 (스트레치)",
}

# ── 색상 팔레트 (양식 표준 기준) ──
PRIMARY_DARK = RGBColor(0x1E, 0x3A, 0x32)         # 진녹 (제목·섹션헤더)
PRIMARY_ACCENT = RGBColor(0x2B, 0x7A, 0x4B)       # 녹 (목표값·강조)
ORANGE_BASELINE = RGBColor(0xF4, 0x9E, 0x33)      # 오렌지 (기준 매출)
WARNING_RED = RGBColor(0xE6, 0x00, 0x00)          # 빨강 (경고 항목)
BODY_GRAY = RGBColor(0x4A, 0x4A, 0x4A)            # 본문 회색
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
STAR_AMBER = RGBColor(0xE6, 0x8A, 0x00)           # ★ 마커 색

# 배경색 (hex string)
KPI_BG = "E8F2EC"
SECTION_HEADER_BG = "1E3A32"
TABLE_HEADER_BG = "1E3A32"
TABLE_AFTER_BG = "2B7A4B"
CATEGORY_BG = "E8F2EC"
ROW_ALT_BG = "F9F9F9"
QUOTE_GRAY_BG = "F2F2F2"
QUOTE_GREEN_BG = "E8F2EC"
KEY_MSG_BG = "EAF0F8"
KEY_MSG_COLOR = RGBColor(0x1B, 0x2D, 0x4E)
KEY_MSG_BORDER_COLOR = "1B2D4E"
BORDER_COLOR = "DDDDDD"
FEE_BORDER_COLOR = "CCCCCC"
FEE_ROW_BG = "E8F2EC"

# ── 폰트 ──
FONT_FAMILY = "Pretendard Variable"
BULLET_FONT = "Cambria Math"

# ── 기호 ──
CIRCLED_NUMBERS = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧"]
BULLET = "▸"
DASH = "–"
STAR = "★"


# ──────────────────────────────────────────
# 유틸리티
# ──────────────────────────────────────────
def _round_to_10k(amount: int) -> int:
    """만원 단위 반올림 (half-up — 5천원 이상은 올림).

    5,787,119 → 5,790,000.
    5,345,000 → 5,350,000 (5천원 이상은 올림).
    5,000     → 10,000    (정확히 5천원도 올림).
    4,999     → 0         (5천원 미만은 내림).
    -1,237,456 → -1,240,000 (부호 유지 후 절대값 반올림).
    0 은 그대로.

    Python 내장 round() 는 banker's rounding(round-half-to-even)이라
    5,345,000 을 5,340,000 으로 떨어뜨리므로 직접 half-up 구현을 사용.

    표시 전용 — 실제 계산값은 원본 JSON 에 유지하고 docx 노출 시점에만 적용.
    """
    if amount == 0:
        return 0
    sign = -1 if amount < 0 else 1
    abs_amt = abs(amount)
    # half-up: (abs + 5000) // 10000 * 10000
    rounded = (abs_amt + 5_000) // 10_000 * 10_000
    return sign * rounded


def _format_won_rounded(amount: int) -> str:
    """원 단위 금액을 만원 반올림 후 '1,234,000원' 형식으로 포맷."""
    return f"{_round_to_10k(amount):,}원"


def _round_down_to_500k(amount: int) -> int:
    """50만원 단위 내림 (버림). 목표 금액 전용.

    보수적 표시로 업주 기대관리에 적합.
    5,787,119 → 5,500,000 / 4,027,952 → 4,000,000
    10만원 미만은 건드리지 않음 (객단가·소액 보호).

    담당자 피드백(2026-04-22): "550만원 정도가 깔끔" —
    목표 금액은 보수적으로 제시해 업주 기대를 관리한다.

    Args:
        amount: 원 단위 금액 (음수/0/10만 미만은 그대로 반환)
    """
    if amount < 100_000:
        return amount
    return (amount // 500_000) * 500_000


def _simplify_target_sub_label(plan: SolutionPlan, raw_sub: str | None) -> str | None:
    """목표매출 KPI 박스의 sub_label 을 한 줄로 간결화.

    기존: "3개월 ₩X → 6개월 ₩Y → 12개월 ₩Z (스트레치) | 현재 × N배 — ... ※담당자 검토 필수..."
    변경: "6개월 목표 ₩Y (+P%)"  (P = 6개월 목표 증가율)

    - tier_plan 이 있으면 tier2_6m.target 사용
    - 없으면 원본 sub_label 유지 (fallback)

    업주 제공용 docx 라 tier 표식(3/6/12M), 로직 설명, 담당자 검토 워닝은 모두 제거.
    """
    tier_plan = plan.tier_plan
    if tier_plan is None:
        return raw_sub

    tier2 = tier_plan.tier2_6m.target
    # 목표 금액은 50만원 단위 내림 (업주 기대관리 — 보수적 표시)
    rounded = _round_down_to_500k(tier2)
    # baseline(현재 월매출) 추출 — core_metrics 의 is_baseline=True 행에서 숫자만 추출
    baseline = 0
    for m in plan.core_metrics:
        if m.is_baseline:
            digits = "".join(ch for ch in m.value if ch.isdigit())
            if digits:
                baseline = int(digits)
            break
    if baseline > 0:
        growth_pct = round((tier2 - baseline) / baseline * 100)
        return f"6개월 목표 ₩{rounded:,} (+{growth_pct}%)"
    return f"6개월 목표 ₩{rounded:,}"


def _reformat_fee_amount(amount: str) -> str:
    """수수료 구조 테이블의 'X,XXX,XXX원' 문자열을 50만원 내림 표기로 재포맷.

    수수료 구조는 목표 금액 구간이므로 50만원 단위 내림 적용.
    '2,880,000원' → '2,500,000원'
    '2,677,752원' → '2,500,000원'
    '4,000,000원' → '4,000,000원' (이미 50만 단위)
    숫자 추출 실패 시 원문 유지.

    담당자 피드백(2026-04-22): 목표 금액은 50만원 단위 내림으로
    보수적 표시 통일.
    """
    digits = "".join(ch for ch in amount if ch.isdigit())
    if not digits:
        return amount
    try:
        val = int(digits)
    except ValueError:
        return amount
    # '원' 포함 여부 유지
    suffix = "원" if "원" in amount else ""
    return f"{_round_down_to_500k(val):,}{suffix}"


def _round_inline_wons(text: str) -> str:
    """본문 문자열 안의 'X,XXX,XXX원' / 'X,XXX원' 패턴을 만원 반올림 표기로 치환.

    key_message·disclaimers 등 업주 눈에 띄는 금액을 담당자 자동 생성 문장 단계에서
    만원 반올림 없이 넣었더라도 docx 노출 시점에 일괄 정리한다.

    5자리 이상(만원 단위 이상)만 대상으로 삼아 주문 건수, 객단가 같은
    "100건" / "12,000원" 수준은 건드리지 않는다 (100,000원 이상부터만 반올림).
    """
    import re
    pattern = re.compile(r"(\d{1,3}(?:,\d{3})+)원")

    def _sub(match: "re.Match[str]") -> str:
        raw = match.group(1)
        digits = raw.replace(",", "")
        if not digits.isdigit():
            return match.group(0)
        val = int(digits)
        # 10만원 미만은 객단가·소액이라 반올림하지 않음
        if val < 100_000:
            return match.group(0)
        return f"{_round_to_10k(val):,}원"

    return pattern.sub(_sub, text)


def _strip_lever_jargon(text: str) -> str:
    """업주 docx 렌더 직전 '레버' 용어 제거·순화 (2026-04-22 담당자 피드백).

    담당자 내부 용어인 '레버'·'4-레버'는 업주에게 의미가 없으므로
    업주 제공용 docx 텍스트 삽입 지점에서 일괄 치환·제거한다.
    내부 산정근거 XLSX(rationale_xlsx) 는 담당자 전용이라 본 함수 적용 대상 아님.

    치환 규칙:
      - "(진짜 레버)" / "( 진짜레버 )" 등 괄호 메타 태그  → 삭제
      - "진짜 레버는 …"                                 → "진짜 성장 동력은 …"
      - "레버 4 (객단가)" / "레버 2 (CTR)" 등           → "{이름}"만 남김
      - "4-레버" / "4레버"                              → "4-지표"
      - 남은 "레버 N" / "레버"(단독)                    → 맥락상 안전히 제거

    공백·구두점 정리 포함.
    """
    if not text:
        return text
    import re

    out = text
    # 1) 괄호 포함 메타 태그 "(진짜 레버)" 완전 삭제
    out = re.sub(r"\s*\(\s*진짜\s*레버\s*\)\s*", "", out)
    # 2) "진짜 레버는" / "진짜 레버가"  → "진짜 성장 동력은/이"
    out = re.sub(r"진짜\s*레버는", "진짜 성장 동력은", out)
    out = re.sub(r"진짜\s*레버가", "진짜 성장 동력이", out)
    out = re.sub(r"진짜\s*레버", "진짜 성장 동력", out)
    # 3) "레버 4 (객단가)" / "레버4(CTR)" 등 → 괄호 안 이름만 남김
    out = re.sub(r"레버\s*\d+\s*\(([^)]+)\)", r"\1", out)
    # 4) "4-레버" / "4 레버" / "4레버" → "4-지표"
    out = re.sub(r"4\s*-?\s*레버", "4-지표", out)
    # 5) 남은 "레버 N" (번호만 붙은 단독 참조) → 삭제
    out = re.sub(r"레버\s*\d+", "", out)
    # 6) 최후의 안전망: 단독 "레버" 단어 제거 (전후 공백 포함)
    out = re.sub(r"\s*레버\s*", " ", out)
    # 7) 공백·구두점 정리
    out = re.sub(r"\s{2,}", " ", out)
    out = out.replace(" ,", ",").replace(" .", ".").replace(" )", ")")
    out = re.sub(r"\(\s+", "(", out)
    out = re.sub(r"\s+\)", ")", out)
    # 빈 괄호 "()"  → 삭제 (치환 잔재)
    out = re.sub(r"\(\s*\)", "", out)
    return out.strip()


def _clean_owner_text(text: str | None) -> str:
    """업주용 텍스트 필터 — 레버 용어 제거 + 공백 정리.

    None 이면 빈 문자열 대신 원본(None) 유지 대신 빈 문자열 반환 금지 —
    빈 문자열을 원래 None 을 받던 옵션 필드로 흘려보내면 렌더가 깨지므로,
    상위 분기에서 None 체크 후에만 호출하는 것을 전제로 한다.
    """
    return _strip_lever_jargon(text or "")


def _set_font(
    run,
    name: str = FONT_FAMILY,
    size_pt: float = 10,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """Run에 폰트를 설정 (한글 eastAsia 포함)."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = r.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _set_cell_bg(cell, hex_color: str) -> None:
    """셀 배경색."""
    tc = cell._element
    tcpr = tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = tc.makeelement(qn("w:shd"), {})
        tcpr.append(shading)
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def _set_cell_borders(
    cell,
    color: str = BORDER_COLOR,
    sz: str = "2",
    sides: list[str] | None = None,
) -> None:
    """셀 테두리."""
    if sides is None:
        sides = ["top", "left", "bottom", "right"]
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side in sides:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def _add_blank_para(doc, space_after_pt: float = 0) -> None:
    """빈 문단 1개 (간격용)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after_pt)


# ──────────────────────────────────────────
# 1. 표지 — 제목 + 부제
# ──────────────────────────────────────────
def _build_title_block(doc: Document, plan: SolutionPlan) -> None:
    """표지 제목·부제 블록."""
    # 제목 + 작성일자 (한 줄)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)

    title_run = para.add_run(plan.document_title)
    _set_font(title_run, size_pt=26, bold=True, color=PRIMARY_DARK)

    if plan.store.document_date:
        date_run = para.add_run(f" ({plan.store.document_date} 기준)")
        _set_font(date_run, size_pt=10, color=BODY_GRAY)

    # 부제: "[업장명] | 배달앱 매출 최적화 컨설팅"
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    sub_run = para.add_run(f"{plan.store.name}  |  {plan.subtitle_suffix}")
    _set_font(sub_run, size_pt=10, color=BODY_GRAY)

    # 위치/업종 (옵션)
    if plan.store.location or plan.store.business_type:
        bits: list[str] = []
        if plan.store.location:
            bits.append(plan.store.location)
        if plan.store.business_type:
            bits.append(plan.store.business_type)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run("  |  ".join(bits))
        _set_font(run, size_pt=9, color=BODY_GRAY)
    else:
        _add_blank_para(doc, space_after_pt=8)


# ──────────────────────────────────────────
# 2. 표지 — 핵심 지표 박스 (가변 N개)
# ──────────────────────────────────────────
def _build_core_metrics(doc: Document, plan: SolutionPlan) -> None:
    """핵심 지표 박스 1행 × N열 (N=2~6)."""
    n = len(plan.core_metrics)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, metric in enumerate(plan.core_metrics):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, KPI_BG)
        cell.text = ""

        # 값 (큰 폰트, 색상 분기)
        if metric.is_baseline:
            value_color = ORANGE_BASELINE
        elif metric.is_target:
            value_color = PRIMARY_ACCENT
        else:
            value_color = PRIMARY_ACCENT

        # 업주 docx 전용 표시 변환:
        #  - 목표 KPI value: 50만원 단위 내림 (담당자 2026-04-22 피드백)
        #  - 기준(baseline) KPI value: 만원 반올림 유지 (현재 매출은 정확도 중요)
        #  - 목표 KPI sub_label: tier/로직/담당자 검토 워닝 제거하고 "6개월 목표 ₩X (+Y%)" 한 줄로 간결화
        #    (객단가·CPC 같은 소액은 건드리지 않음 — 10만원 이상만 반올림 대상)
        display_value = metric.value
        display_sub = metric.sub_label
        if metric.is_target or metric.is_baseline:
            digits = "".join(ch for ch in metric.value if ch.isdigit())
            if digits and "원" in metric.value:
                try:
                    raw_val = int(digits)
                    if raw_val >= 100_000:
                        if metric.is_target:
                            # 목표: 50만원 단위 내림 (보수적 표시)
                            display_value = f"{_round_down_to_500k(raw_val):,}원"
                        else:
                            # 기준(현재): 만원 반올림
                            display_value = f"{_round_to_10k(raw_val):,}원"
                except ValueError:
                    pass
        if metric.is_target:
            display_sub = _simplify_target_sub_label(plan, metric.sub_label)

        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        run = p0.add_run(display_value)
        _set_font(run, size_pt=18, bold=True, color=value_color)

        # 라벨
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        run = p1.add_run(metric.label)
        _set_font(run, size_pt=8.5, color=BODY_GRAY)

        # 보조 라벨 (옵션)
        if display_sub:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(display_sub)
            _set_font(run, size_pt=7.5, color=BODY_GRAY)

    _add_blank_para(doc, space_after_pt=8)


# ──────────────────────────────────────────
# 3. 표지 — 비교표 (가변 행, 멀티라인)
# ──────────────────────────────────────────
def _build_comparison(doc: Document, plan: SolutionPlan) -> None:
    """비교표 (구분 / 지금 / 앞으로)."""
    comp = plan.comparison

    # 제목
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(comp.title)
    _set_font(run, size_pt=14, bold=True, color=PRIMARY_DARK)

    # before/after 라벨 결정
    if comp.header_label == "Before/After":
        before_label, after_label = "Before (현재)", "After (적용 후)"
    else:
        before_label, after_label = "지금", "앞으로"

    n_rows = len(comp.rows) + 1  # 헤더 + 데이터
    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    header_configs = [
        ("구분", TABLE_HEADER_BG),
        (before_label, TABLE_HEADER_BG),
        (after_label, TABLE_AFTER_BG),
    ]
    for i, (text, bg) in enumerate(header_configs):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, bg)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, size_pt=10, bold=True, color=WHITE)

    # 데이터 행
    for row_idx, comp_row in enumerate(comp.rows):
        r = row_idx + 1
        row_bg = ROW_ALT_BG if row_idx % 2 == 0 else "FFFFFF"

        # 구분 셀
        cell = table.cell(r, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, CATEGORY_BG)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(_strip_lever_jargon(comp_row.category))
        _set_font(run, size_pt=10, bold=True, color=PRIMARY_DARK)

        # 지금 (Before) — 멀티라인
        cell = table.cell(r, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        for line_idx, line in enumerate(comp_row.before_lines):
            p = (cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph())
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(_strip_lever_jargon(line))
            _set_font(run, size_pt=9, color=BODY_GRAY)

        # 앞으로 (After) — 멀티라인 (첫 줄 볼드)
        cell = table.cell(r, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        for line_idx, line in enumerate(comp_row.after_lines):
            p = (cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph())
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            is_bold = comp_row.after_first_bold and line_idx == 0
            run = p.add_run(_strip_lever_jargon(line))
            _set_font(
                run,
                size_pt=9,
                bold=is_bold,
                color=PRIMARY_DARK if is_bold else BODY_GRAY,
            )

    _add_blank_para(doc, space_after_pt=4)

    # footer_quote (옵션)
    if comp.footer_quote:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(_strip_lever_jargon(comp.footer_quote))
        _set_font(run, size_pt=9, color=BODY_GRAY)


# ──────────────────────────────────────────
# 3.5. 표지 — 핵심 메시지 박스 (옵션)
# ──────────────────────────────────────────
def _build_key_message(doc: Document, plan: SolutionPlan) -> None:
    """핵심 한줄 메시지 박스 (좌측 두꺼운 border + 연파랑 배경)."""
    if not plan.key_message:
        return

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, KEY_MSG_BG)

    # 좌측 두꺼운 border
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side, sz in [("top", "4"), ("left", "16"), ("bottom", "4"), ("right", "4")]:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), KEY_MSG_BORDER_COLOR)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # 업주 docx: 금액은 만원 반올림 / '레버' 용어는 제거
    run = p.add_run(_strip_lever_jargon(_round_inline_wons(plan.key_message)))
    _set_font(run, size_pt=10, bold=True, color=KEY_MSG_COLOR)

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 3.6. 표지 — 단계별 목표 (3/6/12개월 tier 표)
# ──────────────────────────────────────────
def _build_tier_plan_table(doc: Document, plan: SolutionPlan) -> None:
    """tier_plan(3/6/12개월) 3행 표.

    A/B/C 케이스 전용. tier_plan이 None이면 아무 것도 그리지 않음
    (D/E 케이스 — 단일 목표만 KPI 박스로 표기).

    담당자가 바로 이해할 수 있도록:
      - 헤더: 시점 / 배수 / 목표매출
      - "tier" 영문 금지, "단계별 목표" 한국어 라벨
      - 12개월차는 '스트레치(최종 도달)' 명시
    """
    tier_plan = plan.tier_plan
    if tier_plan is None:
        return

    # 섹션 제목
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("단계별 목표 (3 · 6 · 12개월)")
    _set_font(run, size_pt=12, bold=True, color=PRIMARY_DARK)

    # 표 (헤더 1행 + 데이터 3행) × 3열
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(["시점", "배수", "목표 매출"]):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, TABLE_HEADER_BG)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size_pt=10, bold=True, color=WHITE)

    # 데이터 행 (3/6/12 고정 순서)
    rows_data = [
        ("tier1_3m", tier_plan.tier1_3m),
        ("tier2_6m", tier_plan.tier2_6m),
        ("tier3_12m", tier_plan.tier3_12m),
    ]
    for idx, (key, tier) in enumerate(rows_data):
        r = idx + 1
        row_bg = ROW_ALT_BG if idx % 2 == 0 else "FFFFFF"
        is_stretch = key == "tier3_12m"

        # 시점
        cell = table.cell(r, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, CATEGORY_BG)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(_TIER_LABELS[key])
        _set_font(run, size_pt=10, bold=True, color=PRIMARY_DARK)

        # 배수
        cell = table.cell(r, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"×{tier.multiplier:.2f}")
        _set_font(
            run, size_pt=10, bold=is_stretch,
            color=PRIMARY_ACCENT if is_stretch else PRIMARY_DARK,
        )

        # 목표 매출
        cell = table.cell(r, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{tier.target:,}원")
        _set_font(
            run, size_pt=11, bold=True,
            color=PRIMARY_ACCENT if is_stretch else PRIMARY_DARK,
        )

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 3.7. 표지 — 산정 근거 요약 (목표 매출 rationale)
# ──────────────────────────────────────────
def _build_rationale_section(doc: Document, plan: SolutionPlan) -> None:
    """목표 매출 산정 근거 요약 (회색 박스).

    표시 항목:
      1) 케이스 라벨 (예: "A 케이스 — 성장 잠재력 큼 (스트레치)")
      2) rationale 문장
      3) 상권 TAM 연동 상태 (available=False 면 "미연동" 명시)
      4) 과거 실적 오버라이드 사용 여부

    target_meta/tam_meta 둘 다 None 이면 렌더링 건너뛴다.
    담당자 판독 우선: "TAM/tier" 처음 등장 시 괄호로 한국어 설명.
    """
    meta = plan.target_meta
    tam = plan.tam_meta
    if meta is None and tam is None:
        return

    # 섹션 제목
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("산정 근거 요약")
    _set_font(run, size_pt=12, bold=True, color=PRIMARY_DARK)

    # 근거 문장 생성
    lines: list[str] = []
    if meta is not None:
        if meta.target_case and meta.target_case_label:
            lines.append(
                f"{meta.target_case} 케이스 — {meta.target_case_label}"
            )
        elif meta.target_case_label:
            lines.append(meta.target_case_label)
        if meta.target_rationale:
            lines.append(meta.target_rationale)
        if meta.historical_override_used:
            lines.append(
                "과거 실적 분포(P50/P80) 기반 배수 보정 적용 — "
                "내부 벤치마크로 일반 상수를 대체함"
            )

    # TAM 상태 문구 (available=False 인 경우 반드시 노출)
    if tam is not None:
        if tam.available:
            if tam.tam_monthly_revenue_won:
                lines.append(
                    f"상권 TAM(총시장규모) 연동: 월 약 "
                    f"{tam.tam_monthly_revenue_won:,}원 추정 — 상한 보정 적용"
                )
            else:
                lines.append("상권 TAM(총시장규모) 연동: 상한 보정 적용")
        else:
            reason = f" — {tam.reason}" if tam.reason else ""
            lines.append(
                "상권 TAM(총시장규모) 미연동"
                f"{reason} (현재 × 4.0배 상한 폴백 적용)"
            )

    # 담당자 검토 문구 (자동 산정 한계)
    lines.append(
        "※ 담당자 검토 필수 — 상권·생애주기·경쟁 상황은 자동 산정에 미반영"
    )

    # 회색 박스에 여러 줄 표시
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, QUOTE_GRAY_BG)
    _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        # 경고 문구(마지막 "※")만 이탤릭 색 분기
        is_warning = line.startswith("※")
        run = p.add_run(line)
        _set_font(
            run,
            size_pt=9,
            italic=is_warning,
            bold=(i == 0 and not is_warning),
            color=BODY_GRAY if is_warning else PRIMARY_DARK,
        )

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 3.8. L-2 정본 [목표매출 산정서] — 6개 섹션 빌더
# ──────────────────────────────────────────
def _lr_section_rule(doc: Document, title: str) -> None:
    """정본 §출력 형식의 굵은 구분선(━) + 굵은 섹션 제목."""
    # 상단 구분선
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(6)
    p_top.paragraph_format.space_after = Pt(0)
    r_top = p_top.add_run(_SECTION_RULE_CHAR * _SECTION_RULE_WIDTH)
    _set_font(r_top, size_pt=8, color=PRIMARY_DARK)

    # 제목
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run(title)
    _set_font(r_title, size_pt=12, bold=True, color=PRIMARY_DARK)

    # 하단 구분선
    p_btm = doc.add_paragraph()
    p_btm.paragraph_format.space_before = Pt(0)
    p_btm.paragraph_format.space_after = Pt(4)
    r_btm = p_btm.add_run(_SECTION_RULE_CHAR * _SECTION_RULE_WIDTH)
    _set_font(r_btm, size_pt=8, color=PRIMARY_DARK)


def _lr_gray_box(
    doc: Document,
    lines: list[tuple[str, dict] | str],
    *,
    bg: str = QUOTE_GRAY_BG,
    border_color: str = BORDER_COLOR,
) -> None:
    """회색 박스 — 여러 라인을 각각 폰트 옵션으로 렌더.

    lines 항목은 str(기본 폰트) 또는 (text, font_kwargs) 튜플.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, bg)
    _set_cell_borders(cell, color=border_color, sz="4")
    cell.text = ""
    for i, item in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if isinstance(item, tuple):
            text, kwargs = item
        else:
            text, kwargs = item, {}
        run = p.add_run(text)
        _set_font(
            run,
            size_pt=kwargs.get("size_pt", 9),
            bold=kwargs.get("bold", False),
            italic=kwargs.get("italic", False),
            color=kwargs.get("color", PRIMARY_DARK),
        )
    _add_blank_para(doc, space_after_pt=2)


def _lr_section1_header(doc: Document, plan: SolutionPlan, lr: dict) -> None:
    """섹션 1 — [목표매출 산정서] 헤더 + 현재 매출.

    정본 §출력 형식 라인 190~194:
        매장명 / 카테고리 / 산정일 / 현재 매출 (플랫폼별 + 합계)
    """
    # 최상위 제목
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("[목표매출 산정서]")
    _set_font(r, size_pt=16, bold=True, color=PRIMARY_DARK)

    store = plan.store
    cuisine = lr.get("cuisine", "") or "-"
    doc_date = store.document_date or "-"
    current_baemin = lr.get("analysis", {}).get("current_aov_won")
    # 배민 월 매출은 lever_input.revenue_31d 기반 (tier_plan 에 없으면 KPI 박스에서 추정)
    baemin_revenue_won = None
    # lever_report 상의 revenue_31d 는 직접 저장되어 있지 않으므로 targets.tier_* / growth 역산
    # targets.tier_1 = revenue_31d × (1+imp_s)*(1+ctr_s)*(1+cvr_s)*(1+aov_s)
    targets = lr.get("targets", {})
    tier_1 = targets.get("tier_1_revenue_won", 0)
    growth_1 = targets.get("tier_1_growth_pct", 0.0)
    if growth_1 > -100:
        baemin_revenue_won = int(round(tier_1 / (1 + growth_1 / 100))) if tier_1 else 0

    # 헤더 박스: 매장명/카테고리/산정일/현재 매출
    header_lines: list[tuple[str, dict] | str] = [
        (f"매장명: {store.name}", {"bold": True, "size_pt": 10}),
        (f"카테고리: {cuisine}", {"size_pt": 9}),
        (f"산정일: {doc_date}", {"size_pt": 9}),
    ]
    # 현재 매출 (31일 기준 — 플랫폼별, 만원 반올림)
    if baemin_revenue_won:
        rounded_baemin = _round_to_10k(baemin_revenue_won)
        rev_text = (
            f"현재 매출 (최근 31일 기준): 배민 {rounded_baemin:,}원 / "
            f"쿠팡 데이터 부족 / 요기요 데이터 부족 / 합계 {rounded_baemin:,}원"
        )
    else:
        rev_text = "현재 매출 (최근 31일 기준): 데이터 부족"
    header_lines.append((rev_text, {"bold": True, "color": PRIMARY_ACCENT, "size_pt": 10}))
    # 데이터 부족 표기(빨간 이탤릭)
    header_lines.append(
        ("※ 쿠팡이츠·요기요 데이터 부족 — 추가 수집 필요", {"italic": True, "color": WARNING_RED, "size_pt": 8.5})
    )
    _lr_gray_box(doc, header_lines)


def _lr_section2_levers(doc: Document, lr: dict) -> None:
    """섹션 2 — 레버별 현황 분석 (레버 1~4)."""
    _lr_section_rule(doc, "1. 레버별 현황 분석")

    analysis = lr.get("analysis", {})
    benchmark = analysis.get("cuisine_benchmark", "한식")

    # 정본 §레버 번호·이름 고정 순서
    levers_meta: list[tuple[str, str, str]] = [
        # (타이틀, analysis 필드 key, 단위/표시 문자열 포맷)
        ("레버 1 — 노출수", "impression_delta", "노출수"),
        ("레버 2 — CTR", "ctr_delta", "CTR"),
        ("레버 3 — CVR", "cvr_delta", "CVR"),
        ("레버 4 — 객단가", "aov_delta", "객단가"),
    ]

    # CTR/CVR 벤치 범위는 basis 문자열에 이미 포함되어 있음 → 재가공 없이 노출
    for title, delta_key, _unit in levers_meta:
        delta = analysis.get(delta_key, {})
        short_pct = delta.get("short_term_pct", 0) * 100
        mid_pct = delta.get("mid_term_pct", 0) * 100
        basis = delta.get("basis", "")

        # 레버 헤딩
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(4)
        p_head.paragraph_format.space_after = Pt(2)
        r_head = p_head.add_run(f"[{title}]")
        _set_font(r_head, size_pt=10.5, bold=True, color=PRIMARY_DARK)

        # 현재값·벤치·개선 여력 (회색 박스)
        current_line = _format_current_value(analysis, delta_key, benchmark, lr)
        lines: list[tuple[str, dict] | str] = [
            (f"    현재: {current_line}", {"size_pt": 9}),
            (
                f"    단기 개선 여력: +{short_pct:.0f}% / 근거: {basis}",
                {"size_pt": 9, "color": PRIMARY_DARK},
            ),
            (
                f"    중기 개선 여력: +{mid_pct:.0f}% / 근거: {basis}",
                {"size_pt": 9, "color": PRIMARY_DARK},
            ),
        ]
        _lr_gray_box(doc, lines)


def _format_current_value(
    analysis: dict, delta_key: str, benchmark: str, lr: dict | None = None,
) -> str:
    """레버별 '현재' 라인 포맷. basis 안의 벤치 범위를 그대로 쓰지 않고 현재값만 추림."""
    if delta_key == "impression_delta":
        # L-3: lever_report.current_impressions_31d 가 있으면 실제 값 표기
        imp = (lr or {}).get("current_impressions_31d")
        if isinstance(imp, int) and imp > 0:
            return f"{imp:,}건 (최근 31일)"
        return "노출수 (최근 31일) — 데이터 부족 시 담당자 수동 입력 (정본 §레버1)"
    if delta_key == "ctr_delta":
        ctr = analysis.get("current_ctr_pct", 0)
        return f"{ctr:.2f}% / 벤치({benchmark}) 기준"
    if delta_key == "cvr_delta":
        cvr = analysis.get("current_cvr_pct", 0)
        return f"{cvr:.2f}% / 벤치({benchmark}) 기준"
    if delta_key == "aov_delta":
        aov = analysis.get("current_aov_won", 0)
        return f"{aov:,}원 / {benchmark} 탄력성"
    return "-"


def _lr_section3_targets(doc: Document, lr: dict) -> None:
    """섹션 3 — 목표매출 산정 (1차 3M / 2차 6M). L-3: 플랫폼 분리 필드 활용.

    업주 docx 전용 번호 재매김: §2(레버 분석) 제거로 "1. 목표매출 산정" 이 된다.
    """
    _lr_section_rule(doc, "1. 목표매출 산정")

    targets = lr.get("targets", {})
    tier_1 = targets.get("tier_1_revenue_won", 0)
    tier_2 = targets.get("tier_2_revenue_won", 0)
    growth_1 = targets.get("tier_1_growth_pct", 0.0)
    growth_2 = targets.get("tier_2_growth_pct", 0.0)
    prob_1 = targets.get("tier_1_probability_pct", 0)
    prob_2 = targets.get("tier_2_probability_pct", 0)
    fee_1 = targets.get("tier_1_monthly_fee_won", 0)
    fee_2 = targets.get("tier_2_monthly_fee_won", 0)
    fee_cap_ok = targets.get("fee_cap_ok", True)
    adj_note = targets.get("adjustment_note")
    tam_cap_applied = bool(lr.get("tam_cap_applied", False))

    # L-3: 플랫폼별 baseline 우선 사용 (없으면 tier_1 역산 fallback)
    baemin = targets.get("baemin") or {}
    baemin_baseline = baemin.get("baseline_revenue_won")
    if not baemin_baseline:
        baemin_baseline = (
            int(round(tier_1 / (1 + growth_1 / 100))) if tier_1 and growth_1 > -100 else 0
        )

    platforms_info = {
        "baemin": baemin,
        "coupang_eats": targets.get("coupang_eats") or {},
        "yogiyo": targets.get("yogiyo") or {},
    }

    # 1차 목표 (3개월, 3%) — 녹색 볼드 강조
    _lr_tier_block(
        doc,
        tier_label="[1차 목표 (3% 정산 구간) — 3개월 달성 기준]",
        current_baemin=int(baemin_baseline or 0),
        target_sum=tier_1,
        growth_pct=growth_1,
        prob_pct=prob_1,
        fee_won=fee_1,
        fee_rate_pct=3,
        fee_cap_ok=True,
        adjustment_note=None,
        tam_cap_applied=False,  # 1차엔 TAM 캡 적용 안됨
        platforms=platforms_info,
        tier_key="tier_1",
    )

    # 2차 목표 (6개월, 5%) — 수수료 상한 체크 포함
    _lr_tier_block(
        doc,
        tier_label="[2차 목표 (5% 정산 구간) — 6개월 달성 기준]",
        current_baemin=int(baemin_baseline or 0),
        target_sum=tier_2,
        growth_pct=growth_2,
        prob_pct=prob_2,
        fee_won=fee_2,
        fee_rate_pct=5,
        fee_cap_ok=fee_cap_ok,
        adjustment_note=adj_note,
        tam_cap_applied=tam_cap_applied,
        platforms=platforms_info,
        tier_key="tier_2",
    )


def _lr_tier_block(
    doc: Document,
    *,
    tier_label: str,
    current_baemin: int,
    target_sum: int,
    growth_pct: float,
    prob_pct: int,
    fee_won: int,
    fee_rate_pct: int,
    fee_cap_ok: bool,
    adjustment_note: str | None,
    tam_cap_applied: bool = False,
    platforms: dict[str, dict] | None = None,
    tier_key: str = "tier_1",
) -> None:
    """1차/2차 목표 단일 블록 (녹색 볼드 + 수수료 상한 체크).

    L-3:
      - platforms: {"baemin": PlatformTarget dict, "coupang_eats": ..., "yogiyo": ...}
      - tier_key: "tier_1" 또는 "tier_2" — 각 플랫폼에서 해당 tier 값 조회
      - tam_cap_applied: True 면 "TAM 캡 적용" 경고 라인 추가
    """
    # 라벨
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(tier_label)
    _set_font(r, size_pt=10.5, bold=True, color=PRIMARY_ACCENT)

    platforms = platforms or {}
    target_field = f"{tier_key}_revenue_won"
    growth_field = f"growth_{tier_key[-1]}_pct"  # "tier_1" → "growth_1_pct"

    def _platform_line(name_ko: str, data: dict) -> tuple[str, dict]:
        status = data.get("status", "데이터 부족")
        if status == "산정":
            baseline = int(data.get("baseline_revenue_won") or 0)
            tgt = int(data.get(target_field) or 0)
            gr = float(data.get(growth_field) or 0.0)
            # 현재는 만원 반올림 유지 / 목표는 50만원 내림 (보수적 표시)
            text = (
                f"    {name_ko}: 현재 {_round_to_10k(baseline):,}원 → "
                f"목표 {_round_down_to_500k(tgt):,}원 "
                f"({'+' if gr >= 0 else ''}{gr:.1f}%)"
            )
            return (text, {"size_pt": 9, "bold": True, "color": PRIMARY_ACCENT})
        if status == "현재 유지":
            baseline = int(data.get("baseline_revenue_won") or 0)
            text = (
                f"    {name_ko}: 현재 {_round_to_10k(baseline):,}원 — "
                f"현재 매출 유지 (플랫폼 미운영)"
            )
            return (text, {"size_pt": 9, "italic": True, "color": BODY_GRAY})
        # 데이터 부족
        text = f"    {name_ko}: 데이터 부족 — 산정 제외"
        return (text, {"size_pt": 9, "italic": True, "color": WARNING_RED})

    # 플랫폼별 현재→목표
    lines: list[tuple[str, dict] | str] = [
        _platform_line("배민", platforms.get("baemin", {})),
        _platform_line("쿠팡", platforms.get("coupang_eats", {})),
        _platform_line("요기요", platforms.get("yogiyo", {})),
        (
            # 합계 목표도 50만원 내림 (목표 금액 통일 규칙)
            f"    합계 목표: {_round_down_to_500k(target_sum):,}원 "
            f"(현재 대비 {'+' if growth_pct >= 0 else ''}{growth_pct:.1f}%)",
            {"size_pt": 10, "bold": True, "color": PRIMARY_ACCENT},
        ),
        (f"    달성 확률: {prob_pct}%", {"size_pt": 9}),
        (
            f"    예상 월 수수료: {_round_to_10k(fee_won):,}원 "
            f"(목표 × {fee_rate_pct}%)",
            {"size_pt": 9},
        ),
    ]
    # L-3: TAM 캡 발동 시 명시 (2차만 발동 가능)
    if tam_cap_applied and fee_rate_pct == 5:
        lines.append((
            "    ※ 상권 TAM 점유율 25% 캡 적용됨 — 2차 목표가 상한에서 클리핑됨",
            {"size_pt": 9, "bold": True, "italic": True, "color": WARNING_RED},
        ))
    # 수수료 상한 체크 라인 (2차만 실질 체크. 1차는 항상 적합)
    if fee_rate_pct == 5:
        if fee_cap_ok:
            cap_txt = "    수수료 상한 200만원 체크: 적합"
            lines.append((cap_txt, {"size_pt": 9, "bold": True, "color": PRIMARY_ACCENT}))
        else:
            cap_txt = "    수수료 상한 200만원 체크: 재조정 필요"
            lines.append((cap_txt, {"size_pt": 9, "bold": True, "color": WARNING_RED}))
    _lr_gray_box(doc, lines)

    # adjustment_note (있으면 주황 경고 박스)
    if adjustment_note:
        warn_lines: list[tuple[str, dict] | str] = [
            ("※ 목표 재조정됨", {"bold": True, "size_pt": 9.5, "color": WARNING_RED}),
            (adjustment_note, {"size_pt": 9, "italic": True, "color": PRIMARY_DARK}),
        ]
        _lr_gray_box(doc, warn_lines, bg=WARNING_BG, border_color=WARNING_BORDER)


def _lr_section4_hope_gap(doc: Document, lr: dict, owner_hope_won: int | None) -> None:
    """섹션 4 — 사장님 희망매출 괴리 분석 (희망매출 있을 때만).

    업주 docx 전용 번호 재매김: §2 제거 후 "2. 사장님 희망매출과의 괴리 분석".
    """
    if owner_hope_won is None or owner_hope_won <= 0:
        return

    _lr_section_rule(doc, "2. 사장님 희망매출과의 괴리 분석")

    targets = lr.get("targets", {})
    tier_2 = targets.get("tier_2_revenue_won", 0)
    # 희망매출과의 차이는 표시값(50만 내림) 기준으로 재계산 —
    # 업주 docx 에 표시되는 2차 목표와 diff 가 일치하도록 동기화.
    tier_2_display = _round_down_to_500k(tier_2)
    diff = owner_hope_won - tier_2_display
    diff_pct = (
        round(diff / tier_2_display * 100, 1) if tier_2_display else 0.0
    )

    # 괴리 분석 본문
    analysis = lr.get("analysis", {})
    basis_aov = analysis.get("aov_delta", {}).get("basis", "")

    lines: list[tuple[str, dict] | str] = [
        (
            # 사장님 희망매출 은 입력값(이미 깔끔한 값) — 만원 반올림 유지
            f"사장님 희망매출: {_round_to_10k(owner_hope_won):,}원",
            {"bold": True, "size_pt": 10},
        ),
        (
            # 차이 금액은 표시된 2차 목표(50만 내림)와 산술 일치시킴
            f"2차 목표 대비 차이: {diff:,}원 "
            f"({'+' if diff_pct >= 0 else ''}{diff_pct}%)",
            {"size_pt": 9.5},
        ),
        ("", {}),
        (
            _strip_lever_jargon(
                "괴리 원인: 레버 4 (객단가) 탄력성이 제한적 — "
                f"{basis_aov}"
            ),
            {"size_pt": 9, "color": PRIMARY_DARK},
        ),
        ("", {}),
        ("희망매출 도달 시나리오 (12개월 이상 중장기):", {"bold": True, "size_pt": 9.5}),
        (
            "    ※ 12M 시나리오 확정에는 상권·경쟁·계절성 분석 필요 — 담당자 검토 필수",
            {"italic": True, "color": WARNING_RED, "size_pt": 9},
        ),
        ("    필요 조건 1: (담당자 입력) 카테고리 확장 / 판매채널 추가", {"size_pt": 9}),
        ("    필요 조건 2: (담당자 입력) 배민1 확장 / CPC 상향 운영", {"size_pt": 9}),
        ("    필요 조건 3: (담당자 입력) 쿠팡이츠·요기요 동시 운영", {"size_pt": 9}),
    ]
    _lr_gray_box(doc, lines, bg=WARNING_BG, border_color=WARNING_BORDER)


def _lr_section5_prerequisites(doc: Document, lr: dict) -> None:
    """섹션 5 — 달성 가능성 선행 조건 (업주 제공용: 사장님 실행 필수 항목만).

    업주 docx 전용 정비:
      - 번호 재매김: §2·§6 제거로 "3. 달성 가능성을 좌우하는 선행 조건"
      - "매장주" → "사장님" 전체 치환
      - "제3조 4항 제외 업무" 제거 (계약서 조항 언급 불필요)
      - "기타 / 데이터 추가 수집" 섹션 전체 제거
      - 확률 하락 경고 → 긍정적 동기 문구로 대체
    """
    _lr_section_rule(doc, "3. 달성 가능성을 좌우하는 선행 조건")

    # 사장님 실행 필수 항목
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("[사장님 실행 필수 항목]")
    _set_font(r, size_pt=10, bold=True, color=PRIMARY_DARK)

    owner_lines: list[tuple[str, dict] | str] = [
        (f"    □ {item}", {"size_pt": 9}) for item in _MANDATORY_OWNER_CHECKS
    ]
    _lr_gray_box(doc, owner_lines)

    # 긍정적 동기 문구 (경고 푸터 대체)
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(2)
    p_foot.paragraph_format.space_after = Pt(4)
    r_foot = p_foot.add_run("※ 탄탄한 운영지수는 목표매출 달성의 시작입니다.")
    _set_font(r_foot, size_pt=9, italic=True, color=PRIMARY_DARK)


def _lr_section6_assumptions(doc: Document, lr: dict, plan: SolutionPlan) -> None:
    """섹션 6 — 산정 한계 및 가정 명시."""
    _lr_section_rule(doc, "5. 산정 한계 및 가정 명시")

    analysis = lr.get("analysis", {})
    benchmark = analysis.get("cuisine_benchmark", "-")
    tam_meta = plan.tam_meta

    # 1. 벤치마크 수치
    ctr_pct = analysis.get("current_ctr_pct", 0)
    cvr_pct = analysis.get("current_cvr_pct", 0)
    aov_won = analysis.get("current_aov_won", 0)

    assumption_lines: list[tuple[str, dict] | str] = [
        ("본 산정에 사용된 가정:", {"bold": True, "size_pt": 10}),
        (
            f"    1. 카테고리 벤치마크: {benchmark} "
            f"(CTR {ctr_pct:.2f}% · CVR {cvr_pct:.2f}% · AOV {aov_won:,}원 기준)",
            {"size_pt": 9},
        ),
    ]
    # 2. 상권 경쟁도
    if tam_meta is not None and tam_meta.available:
        if tam_meta.tam_monthly_revenue_won:
            tam_txt = f"연동 (월 TAM 약 {tam_meta.tam_monthly_revenue_won:,}원)"
        else:
            tam_txt = "연동 (상한 보정 적용)"
    elif tam_meta is not None:
        tam_txt = f"미연동 — {tam_meta.reason or '사유 없음'}"
    else:
        tam_txt = "미연동"
    assumption_lines.append(
        (f"    2. 상권 경쟁도: {tam_txt}", {"size_pt": 9})
    )
    # 3. 계절성 — L-3: lever_report.season_factors 를 직접 사용
    season_factors = lr.get("season_factors") or {}
    if season_factors:
        sf_t1 = season_factors.get("tier_1", 1.0)
        sf_t2 = season_factors.get("tier_2", 1.0)
        season_line = (
            f"    3. 계절성 반영 여부: 반영 — tier1 ×{sf_t1:.2f} / tier2 ×{sf_t2:.2f}"
        )
    else:
        # 폴백: tier_plan 의 season_factor 로 추정
        tier_plan = plan.tier_plan
        season_used = False
        if tier_plan is not None:
            season_used = any(
                getattr(tier_plan, f).season_factor != 1.0
                for f in ("tier1_3m", "tier2_6m", "tier3_12m")
            )
        season_line = (
            f"    3. 계절성 반영 여부: {'반영' if season_used else '미반영'}"
        )
    assumption_lines.append((season_line, {"size_pt": 9}))
    # 4. 실행 준수
    assumption_lines.append(
        ("    4. 매장주 실행 준수 가정: 주 1회 이상 협조", {"size_pt": 9})
    )
    # 5. TAM 캡 발동 여부 (L-3)
    tam_cap_applied = bool(lr.get("tam_cap_applied", False))
    tam_cap_line = (
        f"    5. TAM 캡: {'적용 (2차 목표가 상권 TAM×25% 에서 제한됨)' if tam_cap_applied else '미적용'}"
    )
    assumption_lines.append(
        (tam_cap_line, {"size_pt": 9, "bold": tam_cap_applied, "color": (WARNING_RED if tam_cap_applied else None)})
    )
    # 6. 과거 실적 sanity check (L-3)
    hist = lr.get("historical_sanity")
    if hist and hist.get("available"):
        verdict = hist.get("verdict", "")
        n = hist.get("n", 0)
        hist_line = (
            f"    6. 내부 실적 대비 (n={n}): {verdict}"
        )
        is_warn = "공격적" in verdict
        assumption_lines.append(
            (hist_line, {"size_pt": 9, "italic": True,
                         "color": (WARNING_RED if is_warn else BODY_GRAY)})
        )
    elif hist and not hist.get("available"):
        reason = hist.get("reason", "")
        assumption_lines.append(
            (f"    6. 내부 실적 대비: {reason}",
             {"size_pt": 9, "italic": True, "color": BODY_GRAY})
        )

    _lr_gray_box(doc, assumption_lines)

    # 데이터 부족 항목 (disclaimers 재차 요약)
    disclaimers = lr.get("disclaimers", []) or []
    if disclaimers:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(2)
        p_head.paragraph_format.space_after = Pt(2)
        r = p_head.add_run("데이터 부족으로 추정값 적용된 항목:")
        _set_font(r, size_pt=10, bold=True, color=PRIMARY_DARK)
        disc_lines: list[tuple[str, dict] | str] = [
            (f"    - {msg}", {"size_pt": 9, "italic": True, "color": BODY_GRAY})
            for msg in disclaimers
        ]
        _lr_gray_box(doc, disc_lines)


def _build_lever_report_section(plan: SolutionPlan, doc: Document) -> None:
    """lever_report 있을 때 정본 6개 섹션 생성.

    tier 표·근거 박스 대신 사용.
    lever_report=None이면 기존 경로(tier 표 + 근거) 폴백.

    Args:
        plan: SolutionPlan (target_meta.lever_report 유무로 분기)
        doc:  출력 Document
    """
    meta = plan.target_meta
    lr = meta.lever_report if meta else None
    if lr is None:
        # 폴백: 기존 Phase β/γ 경로
        _build_tier_plan_table(doc, plan)
        _build_rationale_section(doc, plan)
        return

    # L-3: owner_hope_won 은 정식 필드 우선, 없으면 disclaimers 파싱 fallback
    owner_hope = lr.get("owner_hope_won")
    if owner_hope is None:
        owner_hope = _extract_owner_hope_from_disclaimers(
            lr.get("disclaimers", []) or []
        )

    # 업주 제공용 섹션 구성:
    #   §2(레버 분석) 과 §6(산정 한계/가정) 은 내부·담당자 용으로 docx 에 노출하지 않는다.
    #   solution_plan.json 에는 lever_report 전문이 유지되므로 담당자는 여전히 조회 가능.
    _lr_section1_header(doc, plan, lr)
    _lr_section3_targets(doc, lr)
    _lr_section4_hope_gap(doc, lr, owner_hope)
    _lr_section5_prerequisites(doc, lr)


def _extract_owner_hope_from_disclaimers(disclaimers: list[str]) -> int | None:
    """disclaimers 문자열에서 사장님 희망매출(원) 추출.

    lever_analysis.build_report() 가 생성하는 포맷:
        "사장님 희망매출 ₩5,000,000원 — 2차 목표 대비 ..."
    """
    import re
    for line in disclaimers:
        # '₩1,234,567' 혹은 '1,234,567원' 패턴
        m = re.search(r"사장님 희망매출[^\d]*([\d,]+)", line)
        if m:
            try:
                return int(m.group(1).replace(",", ""))
            except ValueError:
                continue
    return None


# ──────────────────────────────────────────
# 4. 표지 — 끝 안내문
# ──────────────────────────────────────────
def _build_cover_footer(doc: Document, plan: SolutionPlan) -> None:
    """표지 끝 ※ 안내문."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(plan.cover_footer_note)
    _set_font(run, size_pt=8.5, italic=True, color=BODY_GRAY)


# ──────────────────────────────────────────
# 5. 본문 섹션
# ──────────────────────────────────────────
def _build_section_header(doc: Document, section: Section) -> None:
    """섹션 헤더 (① + 제목, 진녹색 배경)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, SECTION_HEADER_BG)
    cell.text = ""

    idx = section.number - 1
    number = (
        CIRCLED_NUMBERS[idx]
        if 0 <= idx < len(CIRCLED_NUMBERS)
        else f"({section.number})"
    )

    p = cell.paragraphs[0]
    # 업주 docx: '레버' 용어는 섹션 제목에도 노출 금지
    run = p.add_run(f"{number}  {_strip_lever_jargon(section.title)}")
    _set_font(run, size_pt=11, bold=True, color=WHITE)


def _build_section_item(
    doc: Document,
    item: SectionItem,
    serial_num: int | None,
    item_numbering: str,
) -> None:
    """섹션 내 항목 1개."""
    # 항목 제목 (번호 + 마커 + 제목)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(3)

    # 번호 (item_numbering에 따라)
    if item.number:
        num_text = item.number
    elif item_numbering == "serial" and serial_num is not None:
        num_text = str(serial_num)
    else:
        num_text = None

    if num_text:
        num_run = para.add_run(f"{num_text}. ")
        _set_font(num_run, size_pt=11, bold=True, color=PRIMARY_DARK)

    # 마커
    if item.marker == "star":
        m_run = para.add_run(f"{STAR} ")
        _set_font(m_run, size_pt=11, bold=True, color=STAR_AMBER)
    elif item.marker == "warning":
        # 빨간 텍스트로 처리
        pass

    # 제목 (업주 docx: '레버' 용어 제거)
    title_color = WARNING_RED if item.marker == "warning" else PRIMARY_DARK
    title_run = para.add_run(_strip_lever_jargon(item.title))
    _set_font(title_run, size_pt=11, bold=True, color=title_color)

    # bullets (▸)
    for bullet in item.bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 마커
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True, color=PRIMARY_ACCENT)
        # 텍스트 (업주 docx: '레버' 용어 제거)
        t_run = p.add_run(_strip_lever_jargon(bullet))
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)

    # sub_descriptions (–)
    for sub in item.sub_descriptions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        d_run = p.add_run(f"{DASH} ")
        _set_font(d_run, size_pt=9, color=BODY_GRAY)
        t_run = p.add_run(_strip_lever_jargon(sub))
        _set_font(t_run, size_pt=9, color=BODY_GRAY)

    # quote_box (회색/녹색 인용 박스)
    if item.quote_box:
        bg = QUOTE_GREEN_BG if item.quote_box_color == "green" else QUOTE_GRAY_BG
        text_color = PRIMARY_DARK if item.quote_box_color == "green" else BODY_GRAY
        _build_quote_box(doc, _strip_lever_jargon(item.quote_box), bg, text_color)


def _build_quote_box(
    doc: Document, text: str, bg: str, text_color: RGBColor
) -> None:
    """회색/녹색 인용 박스."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, size_pt=9, color=text_color)
    _add_blank_para(doc, space_after_pt=2)


def _build_section(
    doc: Document,
    section: Section,
    item_numbering: str,
    serial_offset: int,
) -> int:
    """섹션 1개 전체 생성. 반환: 직렬 번호 마지막 사용값."""
    _build_section_header(doc, section)
    _add_blank_para(doc, space_after_pt=2)

    serial = serial_offset
    for item_idx, item in enumerate(section.items):
        if item_numbering == "serial":
            serial += 1
            _build_section_item(doc, item, serial_num=serial,
                                item_numbering=item_numbering)
        elif item_numbering == "per_section":
            # 섹션별: section.number.item_idx+1 형태로
            num_str = f"{section.number}.{item_idx + 1}"
            item_with_num = item.model_copy(update={"number": num_str})
            _build_section_item(doc, item_with_num, serial_num=None,
                                item_numbering="per_section")
        else:
            _build_section_item(doc, item, serial_num=None,
                                item_numbering=item_numbering)

    if section.footer_quote:
        _build_quote_box(doc, _strip_lever_jargon(section.footer_quote),
                         QUOTE_GRAY_BG, BODY_GRAY)

    _add_blank_para(doc, space_after_pt=4)
    return serial


# ──────────────────────────────────────────
# 6. 꼬리 — 사장님 확인 (옵션)
# ──────────────────────────────────────────
def _build_owner_checks(
    doc: Document, plan: SolutionPlan, section_number: int
) -> None:
    """⑥ 사장님 확인 필요 사항."""
    if not plan.owner_checks:
        return

    # 가짜 Section으로 헤더 생성
    fake_section = Section(
        number=section_number,
        title="사장님 확인 필요 사항 (VMD 및 메뉴 관련)",
        items=[SectionItem(title="dummy")],
    )
    _build_section_header(doc, fake_section)
    _add_blank_para(doc, space_after_pt=2)

    # 도입 문구 (이탤릭)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(_strip_lever_jargon(plan.owner_check_intro))
    _set_font(run, size_pt=9, italic=True, color=BODY_GRAY)

    # 항목들 (▸ (N) 설명)
    for check in plan.owner_checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True,
                  color=PRIMARY_ACCENT)
        t_run = p.add_run(
            f"({check.number}) {_strip_lever_jargon(check.description)}"
        )
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 7. 꼬리 — 수수료 안내 (옵션)
# ──────────────────────────────────────────
def _build_fee_structure(
    doc: Document, plan: SolutionPlan, section_number: int
) -> None:
    """⑦ 수수료 안내."""
    if not plan.fee_structure:
        return

    fee = plan.fee_structure

    fake_section = Section(
        number=section_number,
        title="수수료 안내",
        items=[SectionItem(title="dummy")],
    )
    _build_section_header(doc, fake_section)
    _add_blank_para(doc, space_after_pt=4)

    # ValueChain 수수료 구조 라벨
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"{plan.company_name} 수수료 구조")
    _set_font(run, size_pt=11, bold=True, color=PRIMARY_DARK)

    # 표 (구분 / 매출 구간 / 수수료율)
    table = doc.add_table(rows=len(fee.tiers) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(["구분", "매출 구간", "수수료율"]):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, TABLE_HEADER_BG)
        _set_cell_borders(cell, color=FEE_BORDER_COLOR, sz="4")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size_pt=9, bold=True, color=WHITE)

    # 구간 행
    for i, tier in enumerate(fee.tiers):
        r = i + 1
        # 라벨
        c = table.cell(r, 0)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, FEE_ROW_BG)
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tier.label)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_DARK)
        # 매출 (만원 반올림 표시)
        c = table.cell(r, 1)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, "FFFFFF")
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ {_reformat_fee_amount(tier.amount)} ]")
        _set_font(run, size_pt=10, bold=True, color=ORANGE_BASELINE)
        # 수수료율
        c = table.cell(r, 2)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, "FFFFFF")
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{tier.rate_pct}%")
        _set_font(run, size_pt=10, bold=True, color=PRIMARY_ACCENT)

    _add_blank_para(doc, space_after_pt=4)

    # 안내 문구
    for note in fee.notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True,
                  color=PRIMARY_ACCENT)
        t_run = p.add_run(_strip_lever_jargon(note))
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)


# ──────────────────────────────────────────
# 8. 푸터 (회사명 + 안내문)
# ──────────────────────────────────────────
def _build_company_footer(doc: Document, plan: SolutionPlan) -> None:
    """문서 끝 푸터."""
    _add_blank_para(doc, space_after_pt=8)

    # 회사명 라인
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{plan.company_name}  |  {plan.company_subtitle}")
    _set_font(run, size_pt=10, color=BODY_GRAY)

    # 안내문
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(plan.company_footer_note)
    _set_font(run, size_pt=8.5, italic=True, color=BODY_GRAY)


# ──────────────────────────────────────────
# 메인 빌더
# ──────────────────────────────────────────
def build_solution_docx(plan: SolutionPlan, output_path: str | Path) -> Path:
    """SolutionPlan v2 → DOCX 결정적 변환."""
    output_path = Path(output_path)
    doc = Document()

    # Normal 스타일
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.15
    style.font.name = FONT_FAMILY
    style.font.size = Pt(10)

    # 페이지 (A4, 여백 2cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))

    # ── 표지 ──
    _build_title_block(doc, plan)
    _build_core_metrics(doc, plan)
    _build_comparison(doc, plan)
    # 담당자 피드백(2026-04-22): 핵심 메시지 박스는 내부 로직 요약이라
    # 업주 제공용 docx 에서 제외. 담당자가 구두로 전달.
    # _build_key_message 함수 자체는 향후 내부용 토글 대비해 유지.
    # _build_key_message(doc, plan)
    # L-2 정본 경로: lever_report 있으면 6-섹션 목표매출 산정서 렌더링.
    # 없으면 기존 Phase β/γ 경로(tier 표 + 근거 박스)로 폴백.
    _build_lever_report_section(plan, doc)
    _build_cover_footer(doc, plan)

    # 페이지 나누기
    doc.add_page_break()

    # ── 본문 섹션 ──
    serial = 0
    for sec in plan.sections:
        serial = _build_section(doc, sec, plan.item_numbering, serial)

    # ── 꼬리 (옵션) ──
    next_section_num = (plan.sections[-1].number if plan.sections else 0) + 1

    if plan.owner_checks:
        _build_owner_checks(doc, plan, next_section_num)
        next_section_num += 1

    if plan.fee_structure:
        _build_fee_structure(doc, plan, next_section_num)

    # ── 푸터 ──
    _build_company_footer(doc, plan)

    # 저장
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
