"""DOCX 빌더 테스트.

Phase β/γ 확장:
  - tier_plan(3/6/12개월) 표 렌더링
  - 산정 근거 섹션 (target_meta + tam_meta)
  - D/E 케이스 단일 목표 fallback
  - 전체 docx 파싱 무결성
"""

import json
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.generator.docx_builder import build_solution_docx  # noqa: E402
from src.schemas.solution import SolutionPlan  # noqa: E402

SAMPLES_DIR = Path(__file__).parent.parent / "data" / "samples"
OUTPUT_DIR = Path(__file__).parent.parent / "output"


def _load(name: str) -> SolutionPlan:
    data = json.loads((SAMPLES_DIR / name).read_text(encoding="utf-8"))
    return SolutionPlan(**data)


def _all_text(doc_path: Path) -> str:
    """docx의 모든 문단 + 표 셀 텍스트를 합쳐 단일 문자열로 반환."""
    doc = Document(str(doc_path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ──────────────────────────────────────────
# 기존 테스트 (선재 실패 — 신 스키마로 복구)
# ──────────────────────────────────────────
def test_build_sample_docx():
    """더피플버거 샘플 데이터(A 케이스)로 DOCX 생성 — tier_plan/tam_meta 포함."""
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_솔루션_계획서.docx"
    result = build_solution_docx(plan, output_path)

    assert result.exists(), f"파일이 생성되지 않았습니다: {result}"
    assert result.stat().st_size > 0, "파일이 비어있습니다"

    # 신 스키마 필드가 실제로 로드됐는지 재확인
    assert plan.tier_plan is not None
    assert plan.target_meta is not None
    assert plan.target_meta.target_case == "A"


# ──────────────────────────────────────────
# L-2: 정본 목표매출 산정서 (lever_report 경로)
# ──────────────────────────────────────────
def test_lever_report_renders_header_section():
    """lever_report 있을 때 [목표매출 산정서] 제목 + 매장명/카테고리/산정일 헤더 포함."""
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_header.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 정본 §출력 형식 최상위 제목
    assert "[목표매출 산정서]" in text
    # 헤더 메타
    assert "매장명: 더피플버거" in text
    assert "카테고리: 한식" in text
    assert "산정일:" in text
    # 현재 매출 — 배민만 있음, 쿠팡/요기요는 "데이터 부족"
    assert "현재 매출" in text
    assert "데이터 부족" in text


def test_lever_report_lever_analysis_section_hidden():
    """업주 docx: §2(레버별 현황 분석) 은 노출하지 않는다.

    (담당자 피드백 2026-04-22) 레버 분석은 solution_plan.json 에는 유지되지만
    업주 제공용 docx 에는 불필요한 내부 로직이라 전체 섹션 생략.
    """
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_levers.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 섹션 제목·레버 블록 모두 미노출
    assert "1. 레버별 현황 분석" not in text
    assert "레버별 현황 분석" not in text
    assert "[레버 1 — 노출수]" not in text
    assert "[레버 2 — CTR]" not in text
    assert "[레버 3 — CVR]" not in text
    assert "[레버 4 — 객단가]" not in text
    # 레버 분석 전용 용어도 미노출
    assert "단기 개선 여력" not in text
    assert "중기 개선 여력" not in text


def test_lever_report_target_calculation_section():
    """섹션 목표매출 산정 — 1차/2차 목표 금액(50만원 내림)·증가율·달성 확률 포함.

    업주 docx 번호 재매김: §2 제거로 "1. 목표매출 산정".
    담당자 2026-04-22 피드백: 목표 금액은 50만원 단위 내림.
    샘플 2,677,752 → 2,500,000 / 3,996,720 → 3,500,000 로 50만원 내림 노출.
    """
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_targets.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 번호 재매김: "1. 목표매출 산정" (기존 "2." 은 제거)
    assert "1. 목표매출 산정" in text
    assert "2. 목표매출 산정" not in text
    assert "[1차 목표 (3% 정산 구간) — 3개월 달성 기준]" in text
    assert "[2차 목표 (5% 정산 구간) — 6개월 달성 기준]" in text
    # 원본 금액(원단위) 및 만원 반올림 값 모두 미노출 — 50만원 내림값만
    assert "2,677,752원" not in text
    assert "3,996,720원" not in text
    assert "2,680,000원" not in text
    # 만원 반올림의 특수 케이스: 3,996,720 → 4,000,000 (만원 반올림) 은 더 이상 등장 안 함
    assert "4,000,000원" not in text
    # 1차 합계 목표 2,677,752 → 2,500,000 (50만 내림)
    assert "2,500,000원" in text
    # 2차 합계 목표 3,996,720 → 3,500,000 (50만 내림)
    assert "3,500,000원" in text
    # 증가율
    assert "+48.8%" in text
    assert "+122.0%" in text
    # 달성 확률
    assert "달성 확률: 75%" in text
    assert "달성 확률: 50%" in text
    # 수수료 상한 적합
    assert "수수료 상한 200만원 체크: 적합" in text


def test_lever_report_no_hope_gap_when_owner_hope_absent():
    """owner_hope 없을 때 괴리 분석 섹션 미생성 — 다음 섹션 번호 재매김 확인."""
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_no_hope.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 샘플 disclaimers 에 '사장님 희망매출' 없음 → 괴리 분석 섹션 생성되지 않음
    assert "사장님 희망매출과의 괴리 분석" not in text
    # 다음 섹션: "3. 달성 가능성을 좌우하는 선행 조건" (§2·§6 제거로 번호 재매김)
    assert "3. 달성 가능성을 좌우하는 선행 조건" in text
    assert "4. 달성 가능성을 좌우하는 선행 조건" not in text


def test_lever_report_hope_gap_when_owner_hope_present():
    """owner_hope 있을 때 괴리 분석 섹션 생성 + 희망 금액 만원 반올림 표시."""
    import copy
    plan = _load("더피플버거_sample.json")
    # lever_report 의 disclaimers 에 희망매출 라인 주입
    lr = copy.deepcopy(plan.target_meta.lever_report)
    lr["disclaimers"].append(
        "사장님 희망매출 ₩5,000,000원 — 2차 목표 대비 +25.1% "
        "(희망치는 괴리 설명용 — 산정 근거 아님)"
    )
    new_meta = plan.target_meta.model_copy(update={"lever_report": lr})
    plan = plan.model_copy(update={"target_meta": new_meta})
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_hope.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 번호 재매김: "2. 사장님 희망매출과의 괴리 분석"
    assert "2. 사장님 희망매출과의 괴리 분석" in text
    assert "3. 사장님 희망매출과의 괴리 분석" not in text
    # 희망 금액(이미 만원 단위)
    assert "5,000,000원" in text
    # 괴리 원인 문구
    assert "괴리 원인" in text
    # 12M 시나리오 placeholder
    assert "담당자 검토 필수" in text


def test_lever_report_prerequisites_tone_reframed():
    """섹션 '선행 조건' 정비 (업주 제공용) — 사장님 호칭·동기 문구·§6 제거 확인."""
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_preq.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)

    # 번호 재매김: "3. 달성 가능성을 좌우하는 선행 조건"
    assert "3. 달성 가능성을 좌우하는 선행 조건" in text
    # "매장주" → "사장님" 전체 치환
    assert "[사장님 실행 필수 항목]" in text
    assert "매장주 실행 필수 항목" not in text
    # 계약서 조항 언급 제거
    assert "제3조 4항" not in text
    assert "제3조 4항 제외 업무" not in text
    # 필수 항목 존재
    assert "메뉴 사진 교체" in text
    assert "리뷰 24시간 내 답변" in text
    assert "영업시간/휴무일 정합성" in text
    # "데이터 추가 수집 필요" 섹션 전체 제거
    assert "데이터 추가 수집" not in text
    # "기타:" 라인 제거 (담당자 협의 사항)
    assert "담당자 협의 사항 기재" not in text
    # 부정적 확률 하락 경고 제거, 긍정 동기 문구로 대체
    assert "1차 목표 달성 확률" not in text
    assert "이하로 하락할 수 있음" not in text
    assert "탄탄한 운영지수는 목표매출 달성의 시작입니다" in text

    # §6(산정 한계 및 가정) 전체 미노출
    assert "산정 한계 및 가정" not in text
    assert "카테고리 벤치마크:" not in text
    assert "상권 경쟁도:" not in text
    assert "계절성 반영 여부:" not in text
    assert "매장주 실행 준수 가정" not in text
    assert "사장님 실행 준수 가정" not in text


def test_lever_report_fee_cap_adjustment_warning():
    """수수료 상한 재조정(adjustment_note) 시 경고 문구 노출."""
    import copy
    plan = _load("더피플버거_sample.json")
    # fee_cap_ok=False + adjustment_note 주입
    lr = copy.deepcopy(plan.target_meta.lever_report)
    lr["targets"]["fee_cap_ok"] = False
    lr["targets"]["adjustment_note"] = (
        "수수료 상한 200만원 적용해 2차 목표를 "
        "₩50,000,000 → ₩40,000,000(4,000만원)으로 재조정"
    )
    new_meta = plan.target_meta.model_copy(update={"lever_report": lr})
    plan = plan.model_copy(update={"target_meta": new_meta})
    output_path = OUTPUT_DIR / f"{plan.store.name}_lever_fee_cap.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 재조정 경고 박스 키워드
    assert "목표 재조정됨" in text
    assert "수수료 상한 200만원 적용해" in text
    assert "수수료 상한 200만원 체크: 재조정 필요" in text


# ──────────────────────────────────────────
# Phase β 회귀: lever_report=None 일 때 기존 tier 표 경로 폴백
# ──────────────────────────────────────────
def test_lever_report_none_falls_back_to_tier_table():
    """lever_report=None 이면 기존 Phase β tier 표 + 근거 박스 경로로 폴백 (회귀 보장)."""
    plan = _load("더피플버거_sample.json")
    # lever_report 제거 (legacy 경로 시뮬레이션)
    new_meta = plan.target_meta.model_copy(update={"lever_report": None})
    plan = plan.model_copy(update={"target_meta": new_meta})
    output_path = OUTPUT_DIR / "lever_none_fallback.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 기존 Phase β tier 표 렌더링 확인
    assert "단계별 목표" in text
    assert "3개월차" in text
    assert "6개월차" in text
    assert "12개월차" in text
    assert "×1.60" in text
    assert "×2.05" in text
    # 기존 rationale 근거 섹션
    assert "산정 근거 요약" in text
    assert "A 케이스" in text
    # 정본 6섹션은 당연히 미생성
    assert "[목표매출 산정서]" not in text
    assert "1. 레버별 현황 분석" not in text


# ──────────────────────────────────────────
# D/E 케이스: tier_plan=None → 단계별 목표 표 없음
# ──────────────────────────────────────────
def test_e_case_no_tier_table():
    """E 케이스(tier_plan=None): 단계별 목표 표 생성 생략, 근거 섹션은 유지."""
    plan = _load("더피플버거_sample_e_case.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_e_case_test.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # tier 관련 라벨이 없어야 함
    assert "단계별 목표" not in text
    assert "3개월차" not in text
    # 그러나 근거 섹션은 있어야 함
    assert "산정 근거 요약" in text
    assert "E 케이스" in text
    assert "신규·소형 매장" in text
    # 단일 목표 매출은 KPI 박스에 남아있어야 함
    assert "3,000,000원" in text


# ──────────────────────────────────────────
# 하위 호환: tier_plan/tam_meta/target_meta 전부 None 인 경우
# ──────────────────────────────────────────
def test_legacy_plan_without_meta_still_builds():
    """Phase β 이전 JSON (메타 필드 없음): 기본 렌더링은 성공, 새 섹션은 생략."""
    plan = _load("더피플버거_sample.json")
    # 강제로 메타 제거 — 구버전 저장본 시뮬레이션
    plan = plan.model_copy(
        update={"tier_plan": None, "tam_meta": None, "target_meta": None}
    )
    output_path = OUTPUT_DIR / "legacy_no_meta_test.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 표지 기본 요소는 유지
    assert "솔루션 계획서" in text
    assert "그래서, 배달 앱 이렇게 바뀝니다" in text
    # 신규 섹션 전부 생략
    assert "단계별 목표" not in text
    assert "산정 근거 요약" not in text


# ──────────────────────────────────────────
# 구조 무결성 — docx 가 python-docx 로 재파싱 가능해야 함
# ──────────────────────────────────────────
@pytest.mark.parametrize(
    "sample_name",
    ["더피플버거_sample.json", "더피플버거_sample_e_case.json"],
)
def test_docx_parses_without_corruption(sample_name: str):
    """생성된 docx 를 다시 python-docx 로 열어도 파싱 오류 없음."""
    plan = _load(sample_name)
    output_path = OUTPUT_DIR / f"{plan.store.name}_parse_test.docx"
    build_solution_docx(plan, output_path)

    doc = Document(str(output_path))
    # 최소 검증: 문단 1개 이상, 표 1개 이상
    assert len(doc.paragraphs) > 5
    assert len(doc.tables) > 0


# ──────────────────────────────────────────
# 2026-04-22 담당자 피드백 반영: 업주 제공용 docx 정비
# ──────────────────────────────────────────
def test_round_to_10k_helper():
    """_round_to_10k: 만원 단위 반올림 — 5천원 이상 올림 / 5천원 미만 내림."""
    from src.generator.docx_builder import _round_to_10k

    # 피드백 명세 예시
    assert _round_to_10k(5_787_119) == 5_790_000   # 2,119 → 내림이 아니라 7,119 이 반올림 → 올림
    assert _round_to_10k(5_345_000) == 5_350_000   # 5천원 이상은 올림
    # 경계 케이스
    assert _round_to_10k(0) == 0
    assert _round_to_10k(4_999) == 0               # 5천원 미만 내림
    assert _round_to_10k(5_000) == 10_000          # 정확히 5천원은 올림
    assert _round_to_10k(14_999) == 10_000         # 4,999 은 내림
    assert _round_to_10k(2_677_752) == 2_680_000   # 샘플 tier_1
    assert _round_to_10k(3_996_720) == 4_000_000   # 샘플 tier_2
    # 음수 (희망매출 괴리 — 희망치가 목표보다 낮을 때)
    assert _round_to_10k(-1_234) == 0
    assert _round_to_10k(-1_237_456) == -1_240_000


def test_round_down_to_500k_helper():
    """_round_down_to_500k: 50만원 단위 내림 (버림). 목표 금액 전용.

    담당자 2026-04-22 피드백: 목표 금액은 50만원 단위로 보수 표시.
    """
    from src.generator.docx_builder import _round_down_to_500k

    # 피드백 명세 예시
    assert _round_down_to_500k(5_787_119) == 5_500_000
    assert _round_down_to_500k(4_027_952) == 4_000_000
    # 경계값
    assert _round_down_to_500k(0) == 0
    assert _round_down_to_500k(99_999) == 99_999       # 10만 미만은 그대로
    assert _round_down_to_500k(100_000) == 0           # 10만(최소 적용) → 50만 미만이라 내림=0
    assert _round_down_to_500k(499_999) == 0           # 50만 미만은 0 으로 내림
    assert _round_down_to_500k(500_000) == 500_000     # 정확히 50만
    assert _round_down_to_500k(999_999) == 500_000     # 100만 미만은 50만으로 내림
    assert _round_down_to_500k(1_000_000) == 1_000_000
    # 샘플 값
    assert _round_down_to_500k(2_677_752) == 2_500_000
    assert _round_down_to_500k(3_996_720) == 3_500_000
    # 큰 값
    assert _round_down_to_500k(100_000_000) == 100_000_000


def test_kpi_target_sub_label_simplified():
    """KPI 박스 '목표 월매출' sub_label 이 한 줄로 간결화돼 노출된다.

    기존: "3개월 ₩... → 6개월 ₩... → 12개월 ₩... | 현재 × N배 — ... ※담당자 검토 필수..."
    변경: "6개월 목표 ₩... (+Y%)"
    """
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_kpi_subsimple.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)

    # 기존 tier 표식·로직·워닝 전부 미노출
    assert "3개월 ₩" not in text
    assert "12개월 ₩" not in text
    assert "(스트레치)" not in text
    assert "4-레버 곱셈 산정" not in text
    assert "담당자 검토 필수: 상권·생애주기" not in text
    assert "담당자 검토 필수: 상권·생애주기·계절" not in text
    # "현재 × N배" 패턴 미노출 (KPI sub_label 에)
    assert "현재 × 2.5배" not in text
    assert "성장 잠재력 큼" not in text

    # 새 sub_label: "6개월 목표 ₩..." 형식 포함
    assert "6개월 목표 ₩" in text


def test_kpi_target_value_floor_to_500k():
    """KPI 박스 '목표 월매출' value 가 50만원 단위 내림으로 표시된다.

    담당자 2026-04-22 피드백: 목표 금액은 50만원 단위 내림으로 보수적 표시.
    5,787,119 → 5,500,000 / 4,027,952 → 4,000,000.
    """
    plan = _load("더피플버거_sample.json")
    # 목표 KPI value 를 원단위로 임의 주입 (피드백 명시 값)
    new_metrics = []
    for m in plan.core_metrics:
        if m.is_target:
            m = m.model_copy(update={"value": "5,787,119원"})
        new_metrics.append(m)
    plan = plan.model_copy(update={"core_metrics": new_metrics})
    output_path = OUTPUT_DIR / f"{plan.store.name}_kpi_value_floor.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 50만원 내림값만 노출 (KPI 박스)
    assert "5,500,000원" in text
    assert "5,787,119원" not in text
    # 만원 반올림 이전 규칙 잔재 없음
    assert "5,790,000원" not in text


def test_fee_amount_floor_to_500k():
    """수수료 구조 테이블의 매출 구간 셀이 50만원 단위 내림으로 재포맷된다.

    담당자 2026-04-22 피드백: 목표 금액은 50만원 단위 내림.
    '2,677,752원' → '2,500,000원'
    '3,996,720원' → '3,500,000원'
    """
    plan = _load("더피플버거_sample.json")
    assert plan.fee_structure is not None
    # 원단위 금액으로 덮어쓰기
    new_tiers = [
        t.model_copy(update={"amount": "2,677,752원"})
        if t.label == "1차 목표 매출"
        else t.model_copy(update={"amount": "3,996,720원"})
        for t in plan.fee_structure.tiers
    ]
    new_fee = plan.fee_structure.model_copy(update={"tiers": new_tiers})
    plan = plan.model_copy(update={"fee_structure": new_fee})
    output_path = OUTPUT_DIR / f"{plan.store.name}_fee_floor.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # 50만원 내림 표시
    assert "2,500,000원" in text
    assert "3,500,000원" in text
    # 원본 원단위 값 미노출
    assert "2,677,752원" not in text
    assert "3,996,720원" not in text
    # 만원 반올림 이전 규칙 잔재 없음
    assert "2,680,000원" not in text
    assert "4,000,000원" not in text


def test_key_message_box_not_present():
    """핵심 메시지 박스가 업주 제공용 docx 에 노출되지 않는다.

    담당자 2026-04-22 피드백: "이것도 삭제. 로직이지 사장님에게 보여줄 게 아님.
    어차피 텍스트로 말해줄테니."
    _build_key_message 함수 자체는 유지되지만 메인 빌더에서 호출하지 않는다.
    """
    plan = _load("더피플버거_sample.json")
    # key_message 가 JSON 에 있더라도 docx 에 나타나지 않아야 함
    assert plan.key_message, "샘플에 key_message 가 있어야 이 테스트가 의미 있음"
    output_path = OUTPUT_DIR / f"{plan.store.name}_no_key_msg.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    # key_message 원문 snippet 이 등장하지 않아야 함
    snippet_head = plan.key_message[:20]
    assert snippet_head not in text, (
        f"핵심 메시지 snippet '{snippet_head}' 가 docx 에 남아있음"
    )
    # 핵심 메시지 특유 시그니처 (샘플 기준 key_message 고유 표현)
    assert "전환율을 끌어올리는 것이 핵심입니다" not in text


def test_final_section_structure_for_owner_docx():
    """업주 제공용 docx 최종 섹션 구조 — 헤더·목표·선행 조건만 포함 / §2·§6 완전 배제."""
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_owner_structure.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)

    # 존재: 헤더·목표매출 산정·선행 조건
    assert "[목표매출 산정서]" in text
    assert "1. 목표매출 산정" in text
    assert "3. 달성 가능성을 좌우하는 선행 조건" in text

    # 미노출: §2 레버 분석 / §6 산정 한계
    assert "레버별 현황 분석" not in text
    assert "산정 한계 및 가정 명시" not in text

    # 내부·담당자 용어 미노출
    assert "매장주" not in text
    assert "제3조 4항" not in text


# ──────────────────────────────────────────
# 2026-04-22 담당자 피드백: 업주 DOCX 에서 '레버' 키워드 전부 제거
# ──────────────────────────────────────────
def test_strip_lever_jargon_parenthetical_tag():
    """'(진짜 레버)' 괄호 메타 태그는 삭제되고 잔여 문자열은 trim 된다."""
    from src.generator.docx_builder import _strip_lever_jargon

    src = "10. 신규 유입 확대 — 광고/노출 강화 (진짜 레버)"
    out = _strip_lever_jargon(src)
    assert "레버" not in out
    assert out == "10. 신규 유입 확대 — 광고/노출 강화"


def test_strip_lever_jargon_numbered_with_name():
    """'레버 4 (객단가)' → '객단가' (번호 제거, 이름만 유지)."""
    from src.generator.docx_builder import _strip_lever_jargon

    src = "괴리 원인: 레버 4 (객단가) 탄력성이 제한적 — 한식 중위 15,500원"
    out = _strip_lever_jargon(src)
    assert "레버" not in out
    # 이름(객단가)은 유지, 탄력성 설명도 보존
    assert "객단가" in out
    assert "탄력성이 제한적" in out
    # 원인 헤더 + 벤치 숫자 보존
    assert out.startswith("괴리 원인:")
    assert "15,500원" in out


def test_strip_lever_jargon_4_lever_multiplication():
    """'4-레버 곱셈' 은 내부 로직 용어 — 잔존 시 '4-지표' 로 순화."""
    from src.generator.docx_builder import _strip_lever_jargon

    assert _strip_lever_jargon("4-레버 곱셈 산정") == "4-지표 곱셈 산정"
    # 하이픈/공백 변형
    assert _strip_lever_jargon("4 레버 곱셈") == "4-지표 곱셈"
    assert _strip_lever_jargon("4레버 산정") == "4-지표 산정"


def test_strip_lever_jargon_real_lever_phrase():
    """'진짜 레버는 …' → '진짜 성장 동력은 …' (의미 보존 순화)."""
    from src.generator.docx_builder import _strip_lever_jargon

    out = _strip_lever_jargon("진짜 레버는 광고 노출 확대입니다")
    assert "레버" not in out
    assert out == "진짜 성장 동력은 광고 노출 확대입니다"


def test_strip_lever_jargon_leaves_non_lever_text_unchanged():
    """'레버' 가 없는 일반 텍스트는 변화가 없다 (양음성 회귀 방지)."""
    from src.generator.docx_builder import _strip_lever_jargon

    samples = [
        "리뷰이벤트 상시 운영 — 재주문 유도",
        "메뉴 사진 교체 (대표 메뉴 최우선)",
        "현재 매출 1,100만원 → 목표 1,800만원",
        "CTR 2.5% / CVR 8.0%",
        "",
    ]
    for s in samples:
        assert _strip_lever_jargon(s) == s


def test_owner_docx_contains_no_lever_keyword_anywhere():
    """업주 제공용 DOCX 를 전체 스캔해 '레버' 단어 0건 보장 (2026-04-22 피드백).

    샘플 SolutionPlan (A 케이스 + lever_report) 로 빌드 →
    python-docx 로 재파싱 → 전체 paragraph + 표 셀 텍스트 합쳐 '레버' substring 스캔.
    """
    plan = _load("더피플버거_sample.json")
    output_path = OUTPUT_DIR / f"{plan.store.name}_no_lever_scan.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    assert "레버" not in text, (
        f"업주 docx 에 '레버' 키워드가 남아있음. "
        f"첫 등장 컨텍스트: ...{text[max(0, text.find('레버') - 40):text.find('레버') + 40]}..."
    )


def test_owner_docx_no_lever_keyword_with_hope_and_fee_adj():
    """owner_hope + 수수료 상한 재조정 경고가 켜진 상태에서도 '레버' 0건."""
    import copy
    plan = _load("더피플버거_sample.json")
    # owner_hope + adjustment_note 동시에 주입 — 최악 시나리오
    lr = copy.deepcopy(plan.target_meta.lever_report)
    lr["disclaimers"].append(
        "사장님 희망매출 ₩5,000,000원 — 2차 목표 대비 +25.1%"
    )
    lr["targets"]["fee_cap_ok"] = False
    lr["targets"]["adjustment_note"] = (
        "수수료 상한 200만원 적용해 2차 목표를 "
        "₩50,000,000 → ₩40,000,000 으로 재조정"
    )
    new_meta = plan.target_meta.model_copy(update={"lever_report": lr})
    plan = plan.model_copy(update={"target_meta": new_meta})

    output_path = OUTPUT_DIR / f"{plan.store.name}_no_lever_scan_max.docx"
    build_solution_docx(plan, output_path)

    text = _all_text(output_path)
    assert "레버" not in text


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
