"""정답 파일 vs 코드 생성물 스타일 diff.

XLSX: 셀단위 값·폰트·배경·병합·크기 비교
DOCX: 문단·표 개수, 폰트 통계, 표 크기 비교
"""
from __future__ import annotations

from collections import defaultdict
from pathlib import Path

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
ANS_XLSX = ROOT / "output/대흥육회_메뉴판_가안.xlsx"
GEN_XLSX = ROOT / "output/roundtrip/대흥육회_메뉴판_가안.xlsx"
ANS_DOCX = ROOT / "output/대흥육회_솔루션_계획서.docx"
GEN_DOCX = ROOT / "output/roundtrip/대흥육회_솔루션_계획서.docx"


def color_rgb(c):
    try:
        return c.rgb if c else None
    except Exception:
        return None


def compare_xlsx():
    print("=" * 60)
    print("[XLSX DIFF] 정답 vs 생성물")
    print("=" * 60)
    wb_a = openpyxl.load_workbook(ANS_XLSX)
    wb_g = openpyxl.load_workbook(GEN_XLSX)

    print(f"시트: 정답={wb_a.sheetnames} / 생성={wb_g.sheetnames}")

    diff_summary: dict[str, int] = defaultdict(int)

    for sname in wb_a.sheetnames:
        if sname not in wb_g.sheetnames:
            print(f"  [{sname}] 생성물에 없음")
            continue
        ws_a = wb_a[sname]
        ws_g = wb_g[sname]
        print(f"\n── [{sname}] ──")
        print(f"  크기: 정답 R{ws_a.max_row}×C{ws_a.max_column} / 생성 R{ws_g.max_row}×C{ws_g.max_column}")

        # 병합
        ma = set(str(r) for r in ws_a.merged_cells.ranges)
        mg = set(str(r) for r in ws_g.merged_cells.ranges)
        missing_merge = ma - mg
        extra_merge = mg - ma
        if missing_merge:
            print(f"  병합 누락: {sorted(missing_merge)}")
            diff_summary["병합 누락"] += len(missing_merge)
        if extra_merge:
            print(f"  병합 추가: {sorted(extra_merge)}")
            diff_summary["병합 추가"] += len(extra_merge)

        # 열너비
        col_diff = []
        for letter in list("ABCDEFGHIJKLMN"):
            wa = ws_a.column_dimensions[letter].width
            wg = ws_g.column_dimensions[letter].width
            if wa and wg and abs(wa - wg) > 0.5:
                col_diff.append(f"{letter}:{wa}→{wg}")
            elif wa and not wg:
                col_diff.append(f"{letter}:{wa}→없음")
            elif wg and not wa:
                col_diff.append(f"{letter}:없음→{wg}")
        if col_diff:
            print(f"  열너비 차이: {col_diff}")
            diff_summary["열너비 차이"] += len(col_diff)

        # 행높이 샘플
        row_diff = []
        for r in range(1, min(ws_a.max_row, ws_g.max_row) + 1):
            ha = ws_a.row_dimensions[r].height
            hg = ws_g.row_dimensions[r].height
            if ha and hg and abs(ha - hg) > 1:
                row_diff.append(f"R{r}:{ha}→{hg}")
        if row_diff[:5]:
            print(f"  행높이 차이 샘플: {row_diff[:5]} (총 {len(row_diff)}행)")
            diff_summary["행높이 차이"] += len(row_diff)

        # 값 비교
        max_r = min(ws_a.max_row, ws_g.max_row)
        max_c = min(ws_a.max_column, ws_g.max_column)
        val_diffs = []
        for r in range(1, max_r + 1):
            for c in range(1, max_c + 1):
                va = ws_a.cell(r, c).value
                vg = ws_g.cell(r, c).value
                if va != vg:
                    val_diffs.append((r, c, va, vg))
        if val_diffs:
            print(f"  값 차이: {len(val_diffs)}개")
            for r, c, va, vg in val_diffs[:5]:
                col_letter = openpyxl.utils.get_column_letter(c)
                print(f"    {col_letter}{r}: {str(va)[:40]!r} → {str(vg)[:40]!r}")
            if len(val_diffs) > 5:
                print(f"    ... +{len(val_diffs) - 5}")
            diff_summary["값 차이"] += len(val_diffs)

        # 스타일 샘플 셀
        sample_cells = [
            ("타이틀", 1, 3),
            ("메뉴헤더", 2, 3),
            ("옵션헤더#", 2, 10),
            ("메뉴그룹", 3, 3),
            ("옵션#박스", 3, 10),
            ("메뉴명", 3, 4),
            ("옵션이름", 3, 11),
            ("옵션항목", 4, 12),
        ]
        for label, r, c in sample_cells:
            if r > max_r or c > max_c:
                continue
            ca = ws_a.cell(r, c)
            cg = ws_g.cell(r, c)
            fa = ca.font
            fg = cg.font
            a_sig = (fa.name, fa.sz, fa.bold, color_rgb(fa.color))
            g_sig = (fg.name, fg.sz, fg.bold, color_rgb(fg.color))
            if a_sig != g_sig:
                print(f"  [{label}@{openpyxl.utils.get_column_letter(c)}{r}] 폰트: {a_sig} → {g_sig}")
                diff_summary["폰트 차이"] += 1
            pa_rgb = color_rgb(ca.fill.start_color) if ca.fill and ca.fill.start_color else None
            pg_rgb = color_rgb(cg.fill.start_color) if cg.fill and cg.fill.start_color else None
            if str(pa_rgb) != str(pg_rgb):
                print(f"  [{label}@{openpyxl.utils.get_column_letter(c)}{r}] 배경: {pa_rgb} → {pg_rgb}")
                diff_summary["배경 차이"] += 1

    print("\n── [XLSX 요약] ──")
    for k, v in diff_summary.items():
        print(f"  {k}: {v}건")


def compare_docx():
    print("\n" + "=" * 60)
    print("[DOCX DIFF] 정답 vs 생성물")
    print("=" * 60)
    doc_a = Document(ANS_DOCX)
    doc_g = Document(GEN_DOCX)

    # 기본 통계
    print(f"문단 수: 정답 {len(doc_a.paragraphs)} / 생성 {len(doc_g.paragraphs)}")
    print(f"표 수:   정답 {len(doc_a.tables)} / 생성 {len(doc_g.tables)}")

    # 각 표 크기
    print("\n── 표별 크기 (행×열) ──")
    for i in range(max(len(doc_a.tables), len(doc_g.tables))):
        a_sz = f"{len(doc_a.tables[i].rows)}×{len(doc_a.tables[i].columns)}" if i < len(doc_a.tables) else "-"
        g_sz = f"{len(doc_g.tables[i].rows)}×{len(doc_g.tables[i].columns)}" if i < len(doc_g.tables) else "-"
        mark = " ←" if a_sz != g_sz else ""
        # 첫 셀 텍스트로 표 식별
        if i < len(doc_a.tables):
            first = doc_a.tables[i].rows[0].cells[0].text[:25].replace("\n", "/")
        elif i < len(doc_g.tables):
            first = doc_g.tables[i].rows[0].cells[0].text[:25].replace("\n", "/")
        else:
            first = ""
        print(f"  표{i + 1} [{first}]: 정답 {a_sz} / 생성 {g_sz}{mark}")

    # 폰트 통계
    def font_stats(doc):
        sizes = defaultdict(int)
        colors = defaultdict(int)
        fonts = defaultdict(int)
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.size:
                    sizes[run.font.size.pt] += 1
                if run.font.color and run.font.color.rgb:
                    colors[str(run.font.color.rgb)] += 1
                if run.font.name:
                    fonts[run.font.name] += 1
        return sizes, colors, fonts

    sa, ca, fa = font_stats(doc_a)
    sg, cg, fg = font_stats(doc_g)
    print(f"\n폰트명 사용: 정답 {dict(fa)} / 생성 {dict(fg)}")
    print(f"폰트 크기 분포: 정답 {dict(sa)} / 생성 {dict(sg)}")
    print(f"색상 분포: 정답 {dict(ca)} / 생성 {dict(cg)}")

    # 문단 내용 sample
    print("\n── 처음 10문단 텍스트 비교 ──")
    for i in range(10):
        ta = doc_a.paragraphs[i].text[:40] if i < len(doc_a.paragraphs) else ""
        tg = doc_g.paragraphs[i].text[:40] if i < len(doc_g.paragraphs) else ""
        mark = " ←" if ta != tg else ""
        print(f"  P{i + 1}: 정답 {ta!r} | 생성 {tg!r}{mark}")


if __name__ == "__main__":
    compare_xlsx()
    compare_docx()
