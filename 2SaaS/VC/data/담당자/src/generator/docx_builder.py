"""솔루션 계획서 DOCX 생성기 v2.

ValueChain 표준 양식 (5건 검증 후 확정)에 맞춘 결정적 빌더.
양식 명세 마스터: data/references/_솔루션양식_표준.md

핵심 설계:
- 가변 KPI 박스 (2~6개)
- 가변 비교표 행 (3~5행, 멀티라인)
- 가변 본문 섹션 (3~8개)
- 마커 분기 (★/▸/빨간경고)
- 회색/녹색 인용 박스
- 꼬리 (사장님 확인 + 수수료) 옵션

동일 입력 = 동일 출력 보장.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.schemas.solution import Section, SectionItem, SolutionPlan

# ── 색상 팔레트 (양식 표준 기준) ──
PRIMARY_DARK = RGBColor(0x1E, 0x3A, 0x32)         # 진녹 (제목·섹션헤더)
PRIMARY_ACCENT = RGBColor(0x2B, 0x7A, 0x4B)       # 녹 (목표값·강조)
ORANGE_BASELINE = RGBColor(0xF4, 0x9E, 0x33)      # 오렌지 (기준 매출)
WARNING_RED = RGBColor(0xE6, 0x00, 0x00)          # 빨강 (경고 항목)
BODY_GRAY = RGBColor(0x4A, 0x4A, 0x4A)            # 본문 회색
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
STAR_AMBER = RGBColor(0xE6, 0x8A, 0x00)           # ★ 마커 색

# 배경색 (hex string)
KPI_BG = "E8F2EC"
SECTION_HEADER_BG = "1E3A32"
TABLE_HEADER_BG = "1E3A32"
TABLE_AFTER_BG = "2B7A4B"
CATEGORY_BG = "E8F2EC"
ROW_ALT_BG = "F9F9F9"
QUOTE_GRAY_BG = "F2F2F2"
QUOTE_GREEN_BG = "E8F2EC"
KEY_MSG_BG = "EAF0F8"
KEY_MSG_COLOR = RGBColor(0x1B, 0x2D, 0x4E)
KEY_MSG_BORDER_COLOR = "1B2D4E"
BORDER_COLOR = "DDDDDD"
FEE_BORDER_COLOR = "CCCCCC"
FEE_ROW_BG = "E8F2EC"

# ── 폰트 ──
FONT_FAMILY = "Pretendard Variable"
BULLET_FONT = "Cambria Math"

# ── 기호 ──
CIRCLED_NUMBERS = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧"]
BULLET = "▸"
DASH = "–"
STAR = "★"


# ──────────────────────────────────────────
# 유틸리티
# ──────────────────────────────────────────
def _set_font(
    run,
    name: str = FONT_FAMILY,
    size_pt: float = 10,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """Run에 폰트를 설정 (한글 eastAsia 포함)."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = r.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _set_cell_bg(cell, hex_color: str) -> None:
    """셀 배경색."""
    tc = cell._element
    tcpr = tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = tc.makeelement(qn("w:shd"), {})
        tcpr.append(shading)
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def _set_cell_borders(
    cell,
    color: str = BORDER_COLOR,
    sz: str = "2",
    sides: list[str] | None = None,
) -> None:
    """셀 테두리."""
    if sides is None:
        sides = ["top", "left", "bottom", "right"]
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side in sides:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def _add_blank_para(doc, space_after_pt: float = 0) -> None:
    """빈 문단 1개 (간격용)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after_pt)


# ──────────────────────────────────────────
# 1. 표지 — 제목 + 부제
# ──────────────────────────────────────────
def _build_title_block(doc: Document, plan: SolutionPlan) -> None:
    """표지 제목·부제 블록."""
    # 제목 + 작성일자 (한 줄)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)

    title_run = para.add_run(plan.document_title)
    _set_font(title_run, size_pt=26, bold=True, color=PRIMARY_DARK)

    if plan.store.document_date:
        date_run = para.add_run(f" ({plan.store.document_date} 기준)")
        _set_font(date_run, size_pt=10, color=BODY_GRAY)

    # 부제: "[업장명] | 배달앱 매출 최적화 컨설팅"
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    sub_run = para.add_run(f"{plan.store.name}  |  {plan.subtitle_suffix}")
    _set_font(sub_run, size_pt=10, color=BODY_GRAY)

    # 위치/업종 (옵션)
    if plan.store.location or plan.store.business_type:
        bits: list[str] = []
        if plan.store.location:
            bits.append(plan.store.location)
        if plan.store.business_type:
            bits.append(plan.store.business_type)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run("  |  ".join(bits))
        _set_font(run, size_pt=9, color=BODY_GRAY)
    else:
        _add_blank_para(doc, space_after_pt=8)


# ──────────────────────────────────────────
# 2. 표지 — 핵심 지표 박스 (가변 N개)
# ──────────────────────────────────────────
def _build_core_metrics(doc: Document, plan: SolutionPlan) -> None:
    """핵심 지표 박스 1행 × N열 (N=2~6)."""
    n = len(plan.core_metrics)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, metric in enumerate(plan.core_metrics):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, KPI_BG)
        cell.text = ""

        # 값 (큰 폰트, 색상 분기)
        if metric.is_baseline:
            value_color = ORANGE_BASELINE
        elif metric.is_target:
            value_color = PRIMARY_ACCENT
        else:
            value_color = PRIMARY_ACCENT

        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        run = p0.add_run(metric.value)
        _set_font(run, size_pt=18, bold=True, color=value_color)

        # 라벨
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        run = p1.add_run(metric.label)
        _set_font(run, size_pt=8.5, color=BODY_GRAY)

        # 보조 라벨 (옵션)
        if metric.sub_label:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(metric.sub_label)
            _set_font(run, size_pt=7.5, color=BODY_GRAY)

    _add_blank_para(doc, space_after_pt=8)


# ──────────────────────────────────────────
# 3. 표지 — 비교표 (가변 행, 멀티라인)
# ──────────────────────────────────────────
def _build_comparison(doc: Document, plan: SolutionPlan) -> None:
    """비교표 (구분 / 지금 / 앞으로)."""
    comp = plan.comparison

    # 제목
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(comp.title)
    _set_font(run, size_pt=14, bold=True, color=PRIMARY_DARK)

    # before/after 라벨 결정
    if comp.header_label == "Before/After":
        before_label, after_label = "Before (현재)", "After (적용 후)"
    else:
        before_label, after_label = "지금", "앞으로"

    n_rows = len(comp.rows) + 1  # 헤더 + 데이터
    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    header_configs = [
        ("구분", TABLE_HEADER_BG),
        (before_label, TABLE_HEADER_BG),
        (after_label, TABLE_AFTER_BG),
    ]
    for i, (text, bg) in enumerate(header_configs):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, bg)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, size_pt=10, bold=True, color=WHITE)

    # 데이터 행
    for row_idx, comp_row in enumerate(comp.rows):
        r = row_idx + 1
        row_bg = ROW_ALT_BG if row_idx % 2 == 0 else "FFFFFF"

        # 구분 셀
        cell = table.cell(r, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, CATEGORY_BG)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(comp_row.category)
        _set_font(run, size_pt=10, bold=True, color=PRIMARY_DARK)

        # 지금 (Before) — 멀티라인
        cell = table.cell(r, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        for line_idx, line in enumerate(comp_row.before_lines):
            p = (cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph())
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            _set_font(run, size_pt=9, color=BODY_GRAY)

        # 앞으로 (After) — 멀티라인 (첫 줄 볼드)
        cell = table.cell(r, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        for line_idx, line in enumerate(comp_row.after_lines):
            p = (cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph())
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            is_bold = comp_row.after_first_bold and line_idx == 0
            run = p.add_run(line)
            _set_font(
                run,
                size_pt=9,
                bold=is_bold,
                color=PRIMARY_DARK if is_bold else BODY_GRAY,
            )

    _add_blank_para(doc, space_after_pt=4)

    # footer_quote (옵션)
    if comp.footer_quote:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(comp.footer_quote)
        _set_font(run, size_pt=9, color=BODY_GRAY)


# ──────────────────────────────────────────
# 3.5. 표지 — 핵심 메시지 박스 (옵션)
# ──────────────────────────────────────────
def _build_key_message(doc: Document, plan: SolutionPlan) -> None:
    """핵심 한줄 메시지 박스 (좌측 두꺼운 border + 연파랑 배경)."""
    if not plan.key_message:
        return

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, KEY_MSG_BG)

    # 좌측 두꺼운 border
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side, sz in [("top", "4"), ("left", "16"), ("bottom", "4"), ("right", "4")]:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), KEY_MSG_BORDER_COLOR)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(plan.key_message)
    _set_font(run, size_pt=10, bold=True, color=KEY_MSG_COLOR)

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 4. 표지 — 끝 안내문
# ──────────────────────────────────────────
def _build_cover_footer(doc: Document, plan: SolutionPlan) -> None:
    """표지 끝 ※ 안내문."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(plan.cover_footer_note)
    _set_font(run, size_pt=8.5, italic=True, color=BODY_GRAY)


# ──────────────────────────────────────────
# 5. 본문 섹션
# ──────────────────────────────────────────
def _build_section_header(doc: Document, section: Section) -> None:
    """섹션 헤더 (① + 제목, 진녹색 배경)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, SECTION_HEADER_BG)
    cell.text = ""

    idx = section.number - 1
    number = (
        CIRCLED_NUMBERS[idx]
        if 0 <= idx < len(CIRCLED_NUMBERS)
        else f"({section.number})"
    )

    p = cell.paragraphs[0]
    run = p.add_run(f"{number}  {section.title}")
    _set_font(run, size_pt=11, bold=True, color=WHITE)


def _build_section_item(
    doc: Document,
    item: SectionItem,
    serial_num: int | None,
    item_numbering: str,
) -> None:
    """섹션 내 항목 1개."""
    # 항목 제목 (번호 + 마커 + 제목)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(3)

    # 번호 (item_numbering에 따라)
    if item.number:
        num_text = item.number
    elif item_numbering == "serial" and serial_num is not None:
        num_text = str(serial_num)
    else:
        num_text = None

    if num_text:
        num_run = para.add_run(f"{num_text}. ")
        _set_font(num_run, size_pt=11, bold=True, color=PRIMARY_DARK)

    # 마커
    if item.marker == "star":
        m_run = para.add_run(f"{STAR} ")
        _set_font(m_run, size_pt=11, bold=True, color=STAR_AMBER)
    elif item.marker == "warning":
        # 빨간 텍스트로 처리
        pass

    # 제목
    title_color = WARNING_RED if item.marker == "warning" else PRIMARY_DARK
    title_run = para.add_run(item.title)
    _set_font(title_run, size_pt=11, bold=True, color=title_color)

    # bullets (▸)
    for bullet in item.bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 마커
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True, color=PRIMARY_ACCENT)
        # 텍스트
        t_run = p.add_run(bullet)
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)

    # sub_descriptions (–)
    for sub in item.sub_descriptions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        d_run = p.add_run(f"{DASH} ")
        _set_font(d_run, size_pt=9, color=BODY_GRAY)
        t_run = p.add_run(sub)
        _set_font(t_run, size_pt=9, color=BODY_GRAY)

    # quote_box (회색/녹색 인용 박스)
    if item.quote_box:
        bg = QUOTE_GREEN_BG if item.quote_box_color == "green" else QUOTE_GRAY_BG
        text_color = PRIMARY_DARK if item.quote_box_color == "green" else BODY_GRAY
        _build_quote_box(doc, item.quote_box, bg, text_color)


def _build_quote_box(
    doc: Document, text: str, bg: str, text_color: RGBColor
) -> None:
    """회색/녹색 인용 박스."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, size_pt=9, color=text_color)
    _add_blank_para(doc, space_after_pt=2)


def _build_section(
    doc: Document,
    section: Section,
    item_numbering: str,
    serial_offset: int,
) -> int:
    """섹션 1개 전체 생성. 반환: 직렬 번호 마지막 사용값."""
    _build_section_header(doc, section)
    _add_blank_para(doc, space_after_pt=2)

    serial = serial_offset
    for item_idx, item in enumerate(section.items):
        if item_numbering == "serial":
            serial += 1
            _build_section_item(doc, item, serial_num=serial,
                                item_numbering=item_numbering)
        elif item_numbering == "per_section":
            # 섹션별: section.number.item_idx+1 형태로
            num_str = f"{section.number}.{item_idx + 1}"
            item_with_num = item.model_copy(update={"number": num_str})
            _build_section_item(doc, item_with_num, serial_num=None,
                                item_numbering="per_section")
        else:
            _build_section_item(doc, item, serial_num=None,
                                item_numbering=item_numbering)

    if section.footer_quote:
        _build_quote_box(doc, section.footer_quote,
                         QUOTE_GRAY_BG, BODY_GRAY)

    _add_blank_para(doc, space_after_pt=4)
    return serial


# ──────────────────────────────────────────
# 6. 꼬리 — 사장님 확인 (옵션)
# ──────────────────────────────────────────
def _build_owner_checks(
    doc: Document, plan: SolutionPlan, section_number: int
) -> None:
    """⑥ 사장님 확인 필요 사항."""
    if not plan.owner_checks:
        return

    # 가짜 Section으로 헤더 생성
    fake_section = Section(
        number=section_number,
        title="사장님 확인 필요 사항 (VMD 및 메뉴 관련)",
        items=[SectionItem(title="dummy")],
    )
    _build_section_header(doc, fake_section)
    _add_blank_para(doc, space_after_pt=2)

    # 도입 문구 (이탤릭)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(plan.owner_check_intro)
    _set_font(run, size_pt=9, italic=True, color=BODY_GRAY)

    # 항목들 (▸ (N) 설명)
    for check in plan.owner_checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True,
                  color=PRIMARY_ACCENT)
        t_run = p.add_run(f"({check.number}) {check.description}")
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)

    _add_blank_para(doc, space_after_pt=4)


# ──────────────────────────────────────────
# 7. 꼬리 — 수수료 안내 (옵션)
# ──────────────────────────────────────────
def _build_fee_structure(
    doc: Document, plan: SolutionPlan, section_number: int
) -> None:
    """⑦ 수수료 안내."""
    if not plan.fee_structure:
        return

    fee = plan.fee_structure

    fake_section = Section(
        number=section_number,
        title="수수료 안내",
        items=[SectionItem(title="dummy")],
    )
    _build_section_header(doc, fake_section)
    _add_blank_para(doc, space_after_pt=4)

    # ValueChain 수수료 구조 라벨
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"{plan.company_name} 수수료 구조")
    _set_font(run, size_pt=11, bold=True, color=PRIMARY_DARK)

    # 표 (구분 / 매출 구간 / 수수료율)
    table = doc.add_table(rows=len(fee.tiers) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(["구분", "매출 구간", "수수료율"]):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, TABLE_HEADER_BG)
        _set_cell_borders(cell, color=FEE_BORDER_COLOR, sz="4")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size_pt=9, bold=True, color=WHITE)

    # 구간 행
    for i, tier in enumerate(fee.tiers):
        r = i + 1
        # 라벨
        c = table.cell(r, 0)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, FEE_ROW_BG)
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tier.label)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_DARK)
        # 매출
        c = table.cell(r, 1)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, "FFFFFF")
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ {tier.amount} ]")
        _set_font(run, size_pt=10, bold=True, color=ORANGE_BASELINE)
        # 수수료율
        c = table.cell(r, 2)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(c, "FFFFFF")
        _set_cell_borders(c, color=FEE_BORDER_COLOR, sz="4")
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{tier.rate_pct}%")
        _set_font(run, size_pt=10, bold=True, color=PRIMARY_ACCENT)

    _add_blank_para(doc, space_after_pt=4)

    # 안내 문구
    for note in fee.notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        b_run = p.add_run(f"{BULLET} ")
        _set_font(b_run, name=BULLET_FONT, size_pt=9, bold=True,
                  color=PRIMARY_ACCENT)
        t_run = p.add_run(note)
        _set_font(t_run, size_pt=9, color=PRIMARY_DARK)


# ──────────────────────────────────────────
# 8. 푸터 (회사명 + 안내문)
# ──────────────────────────────────────────
def _build_company_footer(doc: Document, plan: SolutionPlan) -> None:
    """문서 끝 푸터."""
    _add_blank_para(doc, space_after_pt=8)

    # 회사명 라인
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{plan.company_name}  |  {plan.company_subtitle}")
    _set_font(run, size_pt=10, color=BODY_GRAY)

    # 안내문
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(plan.company_footer_note)
    _set_font(run, size_pt=8.5, italic=True, color=BODY_GRAY)


# ──────────────────────────────────────────
# 메인 빌더
# ──────────────────────────────────────────
def build_solution_docx(plan: SolutionPlan, output_path: str | Path) -> Path:
    """SolutionPlan v2 → DOCX 결정적 변환."""
    output_path = Path(output_path)
    doc = Document()

    # Normal 스타일
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.15
    style.font.name = FONT_FAMILY
    style.font.size = Pt(10)

    # 페이지 (A4, 여백 2cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))

    # ── 표지 ──
    _build_title_block(doc, plan)
    _build_core_metrics(doc, plan)
    _build_comparison(doc, plan)
    _build_key_message(doc, plan)
    _build_cover_footer(doc, plan)

    # 페이지 나누기
    doc.add_page_break()

    # ── 본문 섹션 ──
    serial = 0
    for sec in plan.sections:
        serial = _build_section(doc, sec, plan.item_numbering, serial)

    # ── 꼬리 (옵션) ──
    next_section_num = (plan.sections[-1].number if plan.sections else 0) + 1

    if plan.owner_checks:
        _build_owner_checks(doc, plan, next_section_num)
        next_section_num += 1

    if plan.fee_structure:
        _build_fee_structure(doc, plan, next_section_num)

    # ── 푸터 ──
    _build_company_footer(doc, plan)

    # 저장
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
