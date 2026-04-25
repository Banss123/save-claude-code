"""솔루션 계획서 DOCX 생성기.

SolutionPlan JSON → DOCX 문서를 결정적으로 변환합니다.
동일 입력 = 동일 출력 보장.

원본 기준: 솔루션260327더피플버거.docx (컨설턴트 제작)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.schemas.solution import SolutionPlan

# ── 색상 팔레트 (원본 정밀 추출) ──
PRIMARY_DARK = RGBColor(0x1E, 0x3A, 0x32)
PRIMARY_ACCENT = RGBColor(0x2B, 0x7A, 0x4B)
KPI_BG = "E8F2EC"
SECTION_HEADER_BG = "1E3A32"
TABLE_HEADER_BG = "1E3A32"
TABLE_AFTER_BG = "2B7A4B"
CATEGORY_BG = "E8F2EC"
ROW_ALT_BG = "F9F9F9"
KEY_MSG_BG = "EAF0F8"
KEY_MSG_COLOR = RGBColor(0x1B, 0x2D, 0x4E)
KEY_MSG_BORDER_COLOR = "1B2D4E"
BODY_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
WARNING_RED = RGBColor(0xEE, 0x00, 0x00)  # 섹션 경고 항목용
FEE_RED = RGBColor(0xC0, 0x39, 0x2B)      # 수수료 미발생용
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_COLOR = "DDDDDD"
FEE_BORDER_COLOR = "CCCCCC"
KPI_DIVIDER_COLOR = "AAAAAA"
FEE_ROW_BG = "E8F2EC"

# ── 폰트 설정 ──
FONT_FAMILY = "Pretendard Variable"
BULLET_FONT = "Cambria Math"

# ── 기호 ──
CIRCLED_NUMBERS = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧"]
BULLET = "▸"


def _set_font(run, name: str = FONT_FAMILY, size_pt: float = 10,
              bold: bool = False, color: RGBColor | None = None) -> None:
    """Run에 폰트를 설정합니다."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = r.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _set_cell_bg(cell, hex_color: str) -> None:
    """셀 배경색을 설정합니다."""
    tc = cell._element
    tcpr = tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = tc.makeelement(qn("w:shd"), {})
        tcpr.append(shading)
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def _set_cell_borders(cell, color: str = BORDER_COLOR, sz: str = "2",
                      sides: list[str] | None = None) -> None:
    """셀 테두리를 설정합니다."""
    if sides is None:
        sides = ["top", "left", "bottom", "right"]
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side in sides:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def _build_header(doc: Document, plan: SolutionPlan) -> None:
    """1페이지 상단 헤더 (표지 없이 바로 시작)."""
    # 제목 (space_after=3pt)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run("솔루션 계획서")
    _set_font(run, size_pt=26, bold=True, color=PRIMARY_DARK)

    # 날짜 (space_after=2pt)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"({plan.store.date} 기준)")
    _set_font(run, size_pt=10, color=BODY_GRAY)

    # 업장명 (space_after=10pt → KPI와 간격)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(f"{plan.store.name}  |  배달앱 매출 최적화 컨설팅")
    _set_font(run, size_pt=10, color=BODY_GRAY)


def _build_kpi_table(doc: Document, plan: SolutionPlan) -> None:
    """KPI 4박스 (1행×4열, 값+라벨 별도 paragraph)."""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, kpi in enumerate(plan.kpi_boxes):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, KPI_BG)
        cell.text = ""

        # P0: 값 (20pt, accent, bold, space_after=2pt)
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(2)
        run = p0.add_run(kpi.value)
        _set_font(run, size_pt=20, bold=True, color=PRIMARY_ACCENT)

        # P1: 라벨 (8.5pt, gray)
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(kpi.label)
        _set_font(run, size_pt=8.5, color=BODY_GRAY)

        # 보조 라벨
        if kpi.sub_label:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(kpi.sub_label)
            _set_font(run, size_pt=8, color=BODY_GRAY)

        # 셀간 구분선 (sz=4, 마지막 제외)
        if i < 3:
            _set_cell_borders(cell, color=KPI_DIVIDER_COLOR, sz="4",
                              sides=["right"])

    doc.add_paragraph()


def _build_comparison_table(doc: Document, plan: SolutionPlan) -> None:
    """비교표 (구분/지금/앞으로)."""
    # 제목 (14pt, space_after=6pt)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("그래서, 배달 앱 이렇게 바뀝니다")
    _set_font(run, size_pt=14, bold=True, color=PRIMARY_DARK)

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 (9pt, 테두리 없음)
    header_configs = [
        ("구분", TABLE_HEADER_BG),
        ("지금", TABLE_HEADER_BG),
        ("앞으로", TABLE_AFTER_BG),
    ]
    for i, (text, bg) in enumerate(header_configs):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, bg)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, size_pt=9, bold=True, color=WHITE)

    # 데이터 행 (양쪽 모두 Bold #2B7A4B, 교대 배경, 테두리 있음)
    for row_idx, comp in enumerate(plan.comparison_table):
        row = row_idx + 1
        row_bg = ROW_ALT_BG if row_idx % 2 == 0 else "FFFFFF"

        # 구분
        cell = table.cell(row, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, CATEGORY_BG)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        run = cell.paragraphs[0].add_run(comp.category)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_DARK)

        # 지금 (Bold, accent — 원본과 동일)
        cell = table.cell(row, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        run = cell.paragraphs[0].add_run(comp.before)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_ACCENT)

        # 앞으로
        cell = table.cell(row, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, row_bg)
        _set_cell_borders(cell, color=BORDER_COLOR, sz="2")
        cell.text = ""
        run = cell.paragraphs[0].add_run(comp.after)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_ACCENT)

    doc.add_paragraph()


def _build_key_message(doc: Document, plan: SolutionPlan) -> None:
    """핵심 메시지 박스 (4면 border: top=4,left=16,bottom=4,right=4)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, KEY_MSG_BG)

    # 4면 border (좌측만 두껍게)
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side, sz in [("top", "4"), ("left", "16"), ("bottom", "4"), ("right", "4")]:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), KEY_MSG_BORDER_COLOR)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(plan.key_message)
    _set_font(run, size_pt=9, bold=False, color=KEY_MSG_COLOR)

    doc.add_paragraph()


def _build_mid_message(doc: Document, text: str) -> None:
    """중간 메시지 박스 (섹션 사이 강조, key_message와 동일 스타일)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, KEY_MSG_BG)

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = tc.makeelement(qn("w:tcBorders"), {})
    for side, sz in [("top", "4"), ("left", "16"), ("bottom", "4"), ("right", "4")]:
        el = tc.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), KEY_MSG_BORDER_COLOR)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_font(run, size_pt=9, bold=False, color=KEY_MSG_COLOR)


def _build_guide_text(doc: Document) -> None:
    """안내 문구 (8.5pt, italic, space_after=10pt)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(
        "※ 아래 세부 계획을 검토하신 후, "
        "조정이 필요한 부분을 말씀해주시면 반영하겠습니다."
    )
    _set_font(run, size_pt=8.5, color=BODY_GRAY)
    run.font.italic = True


def _build_section(doc: Document, section, index: int) -> None:
    """섹션 1개를 생성합니다."""
    # 섹션 헤더 (10pt, 원본 기준 Normal 상속)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(header_cell, SECTION_HEADER_BG)
    header_cell.text = ""

    number = (CIRCLED_NUMBERS[index]
              if index < len(CIRCLED_NUMBERS) else f"({index + 1})")
    p = header_cell.paragraphs[0]
    # 원본: "①  배민 기본 세팅" (더블 스페이스)
    run = p.add_run(f"{number}  {section.title}")
    _set_font(run, size_pt=10, bold=True, color=WHITE)

    # 항목들
    for item_idx, item in enumerate(section.items, 1):
        # 항목 제목 (11pt, space_before=10pt, space_after=3pt)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(3)

        # 번호 (accent green, 분리 run)
        num_run = para.add_run(f"{item_idx}. ")
        _set_font(num_run, size_pt=11, bold=True, color=PRIMARY_ACCENT)

        # 제목 (경고: #EE0000 12pt / 일반: #1E3A32 11pt)
        if item.is_warning:
            title_run = para.add_run(item.title)
            _set_font(title_run, size_pt=12, bold=True, color=WARNING_RED)
        else:
            title_run = para.add_run(item.title)
            _set_font(title_run, size_pt=11, bold=True, color=PRIMARY_DARK)

        # 설명 (▸ + 3 run 분리: ▸ / key phrase bold accent / → rest normal)
        if item.description:
            desc_para = doc.add_paragraph()
            desc_para.paragraph_format.space_before = Pt(1.5)
            desc_para.paragraph_format.space_after = Pt(1.5)

            # Run 1: ▸ (Cambria Math, bold, accent)
            br = desc_para.add_run(f"{BULLET} ")
            _set_font(br, name=BULLET_FONT, size_pt=9,
                      bold=True, color=PRIMARY_ACCENT)

            # → 기준 분리
            if "\u2192" in item.description or "→" in item.description:
                sep = "→" if "→" in item.description else "\u2192"
                parts = item.description.split(sep, 1)
                # Run 2: key phrase (bold, accent)
                kr = desc_para.add_run(parts[0].rstrip())
                _set_font(kr, size_pt=9, bold=True, color=PRIMARY_ACCENT)
                # Run 3: → rest (normal, inherit color)
                rr = desc_para.add_run(f" {sep} {parts[1].lstrip()}")
                _set_font(rr, size_pt=9)
            else:
                # 전체 bold accent
                kr = desc_para.add_run(item.description)
                _set_font(kr, size_pt=9, bold=True, color=PRIMARY_ACCENT)

        # 하위 항목 (▸, 1.5pt spacing)
        for sub in item.sub_items:
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.space_before = Pt(1.5)
            sub_para.paragraph_format.space_after = Pt(1.5)

            br = sub_para.add_run(f"{BULLET} ")
            _set_font(br, name=BULLET_FONT, size_pt=9,
                      bold=True, color=PRIMARY_ACCENT)

            sr = sub_para.add_run(sub)
            _set_font(sr, size_pt=9, bold=True, color=PRIMARY_ACCENT)

    # 섹션 간 빈 행 (space_before=4pt)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)


def _build_fee_table(doc: Document, plan: SolutionPlan) -> None:
    """수수료 구조 테이블 (border sz=4, color=CCCCCC)."""
    if not plan.fee_structure:
        return

    fee = plan.fee_structure
    rows = len(fee.tiers) + 2  # 헤더 + 구간 + 미발생 조건
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 ("매출 구간" / "수수료율")
    for i, h in enumerate(["매출 구간", "수수료율"]):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, TABLE_HEADER_BG)
        _set_cell_borders(cell, color=FEE_BORDER_COLOR, sz="4")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_font(run, size_pt=9, bold=True, color=WHITE)

    # 구간 행
    for i, tier in enumerate(fee.tiers):
        row = i + 1
        cond_cell = table.cell(row, 0)
        cond_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cond_cell, FEE_ROW_BG)
        _set_cell_borders(cond_cell, color=FEE_BORDER_COLOR, sz="4")
        cond_cell.text = ""
        run = cond_cell.paragraphs[0].add_run(tier.condition)
        _set_font(run, size_pt=9)

        rate_cell = table.cell(row, 1)
        rate_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(rate_cell, "FFFFFF")
        _set_cell_borders(rate_cell, color=FEE_BORDER_COLOR, sz="4")
        rate_cell.text = ""
        run = rate_cell.paragraphs[0].add_run(tier.rate)
        _set_font(run, size_pt=9, bold=True, color=PRIMARY_ACCENT)

    # 미발생 조건 (2열 유지)
    last = rows - 1
    cond_cell = table.cell(last, 0)
    cond_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cond_cell, FEE_ROW_BG)
    _set_cell_borders(cond_cell, color=FEE_BORDER_COLOR, sz="4")
    cond_cell.text = ""
    run = cond_cell.paragraphs[0].add_run(fee.no_fee_condition)
    _set_font(run, size_pt=9)

    rate_cell = table.cell(last, 1)
    rate_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(rate_cell, "FFFFFF")
    _set_cell_borders(rate_cell, color=FEE_BORDER_COLOR, sz="4")
    rate_cell.text = ""
    run = rate_cell.paragraphs[0].add_run("수수료 미발생")
    _set_font(run, size_pt=9, bold=True, color=FEE_RED)


def build_solution_docx(plan: SolutionPlan, output_path: str | Path) -> Path:
    """SolutionPlan → DOCX 문서를 생성합니다."""
    output_path = Path(output_path)
    doc = Document()

    # Normal 스타일 (원본 기준: 맑은 고딕, 10pt, line_spacing 1.15)
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.15
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # 페이지 설정 (A4, 여백 2cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))

    # 1. 헤더 (표지 없이 1페이지 바로 시작)
    _build_header(doc, plan)

    # 2. KPI 박스
    _build_kpi_table(doc, plan)

    # 3. 비교표
    _build_comparison_table(doc, plan)

    # 4. 핵심 메시지
    _build_key_message(doc, plan)

    # 5. 안내 문구
    _build_guide_text(doc)

    # 페이지 나누기
    doc.add_page_break()

    # 6. 섹션들 (+ 중간 메시지 박스)
    for i, sec in enumerate(plan.sections):
        _build_section(doc, sec, i)
        if sec.message:
            _build_mid_message(doc, sec.message)

    # 7. 수수료 테이블
    _build_fee_table(doc, plan)

    # 저장
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
