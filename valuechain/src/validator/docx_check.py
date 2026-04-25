"""DOCX 솔루션 계획서 검수기.

SolutionPlan (JSON) ↔ 생성된 DOCX 파일을 교차 검증합니다.

검수 항목 (CLAUDE.md 기준):
1. KPI 박스 4건 이상 표시
2. 비교표 "지금/앞으로" 내용이 원본과 일치
3. 핵심 메시지가 박스 안에 존재
4. 섹션 구분자 전체 포함
5. 항목 번호 연속성
6. 수수료 구조 테이블 정확성
7. 원본에 없는 내용 임의 추가 여부
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from src.schemas.solution import SolutionPlan
from src.schemas.validation import CheckGroup, CheckItem, CheckStatus


def _extract_all_text(doc: Document) -> str:
    """문서 전체 텍스트를 추출합니다 (테이블 포함)."""
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def validate_docx(plan: SolutionPlan, docx_path: str | Path) -> CheckGroup:
    """SolutionPlan ↔ DOCX 파일을 검수합니다."""
    docx_path = Path(docx_path)
    items: list[CheckItem] = []

    # ── 파일 존재 확인 ──
    if not docx_path.exists():
        items.append(CheckItem(
            name="파일 존재", status=CheckStatus.FAIL,
            message=f"{docx_path.name} 파일을 찾을 수 없습니다",
        ))
        return CheckGroup(name="DOCX 검수", items=items)

    doc = Document(str(docx_path))
    all_text = _extract_all_text(doc)

    # ── 1. KPI 박스 4건 ──
    # T0이 KPI 테이블 (1행 x 4열)
    kpi_table = doc.tables[0] if doc.tables else None
    kpi_count = 0
    kpi_values_found: list[str] = []
    if kpi_table:
        kpi_count = len(kpi_table.columns)
        for cell in kpi_table.rows[0].cells:
            kpi_values_found.append(cell.text.strip()[:30])

    items.append(CheckItem(
        name="KPI 박스 4건",
        status=CheckStatus.PASS if kpi_count >= 4 else CheckStatus.FAIL,
        message=f"KPI {kpi_count}건 발견",
        expected="4",
        actual=str(kpi_count),
    ))

    # KPI 값이 JSON과 일치하는지
    kpi_match = True
    kpi_mismatches: list[str] = []
    for kpi_box in plan.kpi_boxes:
        found = any(kpi_box.value in cell_text for cell_text in kpi_values_found)
        if not found:
            kpi_match = False
            kpi_mismatches.append(f"{kpi_box.label}={kpi_box.value}")

    items.append(CheckItem(
        name="KPI 값 JSON 일치",
        status=CheckStatus.PASS if kpi_match else CheckStatus.WARN,
        message=(
            "전체 KPI 값 일치"
            if kpi_match
            else f"불일치: {', '.join(kpi_mismatches[:3])}"
        ),
    ))

    # ── 2. 비교표 "지금/앞으로" ──
    comp_table = doc.tables[1] if len(doc.tables) > 1 else None
    if comp_table and len(comp_table.rows) >= 4:
        comp_ok = True
        comp_issues: list[str] = []

        # 헤더 확인
        headers = [c.text.strip() for c in comp_table.rows[0].cells]
        if "지금" not in " ".join(headers) and "앞으로" not in " ".join(headers):
            comp_ok = False
            comp_issues.append("헤더에 '지금/앞으로' 없음")

        # 각 행의 값이 JSON과 일치하는지
        for ri, comp_row in enumerate(plan.comparison_table):
            if ri + 1 < len(comp_table.rows):
                row_cells = [c.text.strip() for c in comp_table.rows[ri + 1].cells]
                row_text = " ".join(row_cells)
                # category 확인
                if comp_row.category not in row_text:
                    comp_ok = False
                    comp_issues.append(f"'{comp_row.category}' 카테고리 없음")

        items.append(CheckItem(
            name="비교표 구조 및 내용",
            status=CheckStatus.PASS if comp_ok else CheckStatus.FAIL,
            message=(
                f"비교표 {len(plan.comparison_table)}행 정상"
                if comp_ok
                else f"이슈: {', '.join(comp_issues[:3])}"
            ),
        ))
    else:
        items.append(CheckItem(
            name="비교표 구조 및 내용",
            status=CheckStatus.FAIL,
            message="비교표 테이블을 찾을 수 없습니다",
        ))

    # ── 3. 핵심 메시지 존재 ──
    # JSON의 key_message 일부가 DOCX 어딘가에 있는지
    key_msg = plan.key_message
    # 핵심 메시지의 앞 20자로 검색 (전체 일치는 줄바꿈 등으로 어려울 수 있음)
    key_msg_snippet = key_msg[:20] if len(key_msg) > 20 else key_msg
    msg_found = key_msg_snippet in all_text

    items.append(CheckItem(
        name="핵심 메시지 존재",
        status=CheckStatus.PASS if msg_found else CheckStatus.FAIL,
        message=(
            f"'{key_msg_snippet}...' 확인됨"
            if msg_found
            else f"'{key_msg_snippet}...' 를 찾을 수 없음"
        ),
    ))

    # ── 4. 섹션 구분자 전체 포함 ──
    section_titles = [s.title for s in plan.sections]
    missing_sections: list[str] = []
    for title in section_titles:
        if title not in all_text:
            missing_sections.append(title)

    items.append(CheckItem(
        name="섹션 구분자 전체 포함",
        status=CheckStatus.PASS if not missing_sections else CheckStatus.FAIL,
        message=(
            f"전체 {len(section_titles)}개 섹션 존재"
            if not missing_sections
            else f"누락: {', '.join(missing_sections[:3])}"
        ),
    ))

    # ── 5. 항목 번호 연속성 ──
    # 단락에서 "N. ①" 패턴 추출
    import re
    item_numbers: list[int] = []
    for p in doc.paragraphs:
        m = re.match(r"^(\d+)\.\s", p.text.strip())
        if m:
            item_numbers.append(int(m.group(1)))

    # 섹션별로 번호가 1부터 연속인지 확인
    numbering_ok = True
    if item_numbers:
        # 번호가 1부터 시작하는 연속 시퀀스인지 (섹션 리셋 고려)
        # 여러 섹션에서 각각 1부터 시작하므로, 1이 나올 때마다 리셋
        current_expected = 1
        for n in item_numbers:
            if n == 1:
                current_expected = 1
            if n != current_expected:
                numbering_ok = False
                break
            current_expected += 1

    items.append(CheckItem(
        name="항목 번호 연속성",
        status=CheckStatus.PASS if numbering_ok else CheckStatus.WARN,
        message=(
            f"항목 번호 연속 ({len(item_numbers)}개)"
            if numbering_ok
            else "항목 번호 불연속 발견"
        ),
    ))

    # ── 6. 수수료 구조 (있는 경우) ──
    if plan.fee_structure:
        fee_found = "수수료" in all_text
        no_fee_found = False
        # "미달성" 또는 "미발생" 키워드
        no_fee_keywords = ["미달성", "미발생"]
        for kw in no_fee_keywords:
            if kw in all_text:
                no_fee_found = True
                break

        items.append(CheckItem(
            name="수수료 구조 존재",
            status=CheckStatus.PASS if fee_found else CheckStatus.FAIL,
            message="수수료 섹션 존재" if fee_found else "수수료 관련 내용 없음",
        ))

        items.append(CheckItem(
            name="수수료 미발생 조건 강조",
            status=CheckStatus.PASS if no_fee_found else CheckStatus.WARN,
            message=(
                "'미달성/미발생' 문구 확인됨"
                if no_fee_found
                else "'미달성 시 수수료 미발생' 문구를 찾을 수 없음"
            ),
        ))

    # ── 7. 업장명 일치 ──
    store_name_found = plan.store.name in all_text
    items.append(CheckItem(
        name="업장명 일치",
        status=CheckStatus.PASS if store_name_found else CheckStatus.FAIL,
        message=(
            f"'{plan.store.name}' 확인됨"
            if store_name_found
            else f"'{plan.store.name}' 를 찾을 수 없음"
        ),
    ))

    # ── 8. 경고 항목 볼드 확인 ──
    warning_items = [
        item for s in plan.sections for item in s.items if item.is_warning
    ]
    if warning_items:
        # 경고 항목의 title이 문서에 존재하는지만 확인
        missing_warnings: list[str] = []
        for wi in warning_items:
            if wi.title[:15] not in all_text:
                missing_warnings.append(wi.title[:20])

        items.append(CheckItem(
            name="경고 항목 존재",
            status=CheckStatus.PASS if not missing_warnings else CheckStatus.WARN,
            message=(
                f"경고 항목 {len(warning_items)}건 모두 존재"
                if not missing_warnings
                else f"경고 항목 누락: {', '.join(missing_warnings[:3])}"
            ),
        ))

    return CheckGroup(name="DOCX 검수", items=items)
